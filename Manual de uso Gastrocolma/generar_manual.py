# -*- coding: utf-8 -*-
"""
Script para generar el Manual de Usuario de la Aplicación GPP
Genera imágenes ilustrativas y documento Word completo
"""

import os
import sys
import io

# Fix encoding for Windows console
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime

# Ruta de salida
OUTPUT_DIR = r"C:\Users\ASUS\OneDrive\Desktop\Manual de uso Gastrocolma"
IMG_DIR = os.path.join(OUTPUT_DIR, "imagenes")
os.makedirs(IMG_DIR, exist_ok=True)

# ============================================================
# PARTE 1: GENERACIÓN DE IMÁGENES ILUSTRATIVAS
# ============================================================

def get_font(size=14, bold=False):
    """Obtiene fuente del sistema"""
    try:
        if bold:
            return ImageFont.truetype("arialbd.ttf", size)
        return ImageFont.truetype("arial.ttf", size)
    except:
        try:
            if bold:
                return ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", size)
            return ImageFont.truetype("C:/Windows/Fonts/arial.ttf", size)
        except:
            return ImageFont.load_default()


def crear_imagen_bienvenida():
    """Crea imagen ilustrativa de la pantalla de inicio"""
    fig, ax = plt.subplots(1, 1, figsize=(12, 7))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis('off')
    fig.patch.set_facecolor('#F8FAFC')

    # Sidebar simulado
    sidebar = FancyBboxPatch((0, 0), 2.8, 7, boxstyle="round,pad=0.05",
                              facecolor='#172139', edgecolor='none')
    ax.add_patch(sidebar)

    ax.text(1.4, 6.5, "Navegacion", fontsize=13, fontweight='bold',
            color='white', ha='center', va='center')

    menu_items = ["Inicio", "Proceso\nAdministrativo", "Proceso\nOperativo",
                  "Resultados\nGenerales", "Prioridades", "Guardar/Cargar"]
    icons = ["[1]", "[2]", "[3]", "[4]", "[5]", "[6]"]
    y_pos = [5.8, 5.1, 4.4, 3.7, 3.0, 2.3]

    for i, (item, icon, y) in enumerate(zip(menu_items, icons, y_pos)):
        color = '#00AEAC' if i == 0 else '#2A3A5C'
        rect = FancyBboxPatch((0.2, y - 0.25), 2.4, 0.5, boxstyle="round,pad=0.05",
                               facecolor=color, edgecolor='none', alpha=0.8)
        ax.add_patch(rect)
        ax.text(0.5, y, icon, fontsize=8, color='#00AEAC', ha='left', va='center')
        ax.text(1.0, y, item, fontsize=8, color='white', ha='left', va='center')

    main_bg = FancyBboxPatch((3.1, 0.2), 8.7, 6.6, boxstyle="round,pad=0.1",
                              facecolor='white', edgecolor='#E2E8F0')
    ax.add_patch(main_bg)

    banner = FancyBboxPatch((3.4, 4.5), 8.1, 2.0, boxstyle="round,pad=0.1",
                             facecolor='#172139', edgecolor='none')
    ax.add_patch(banner)

    for i in range(20):
        x = 3.4 + (8.1 * i / 20)
        color_val = 0.09 + (i / 20) * 0.1
        rect = plt.Rectangle((x, 4.5), 8.1/20, 2.0, color=(color_val, 0.34 + i*0.01, 0.34 + i*0.01), alpha=0.5)
        ax.add_patch(rect)

    ax.text(7.45, 5.5, "Bienvenido al Sistema GPP", fontsize=18, fontweight='bold',
            color='white', ha='center', va='center')
    ax.text(7.45, 4.9, "Sistema Profesional de Evaluacion", fontsize=11,
            color='#CCCCCC', ha='center', va='center')

    card1 = FancyBboxPatch((3.4, 1.5), 4.8, 2.7, boxstyle="round,pad=0.1",
                            facecolor='white', edgecolor='#E2E8F0')
    ax.add_patch(card1)
    ax.text(3.7, 3.9, "Guia de Uso", fontsize=12, fontweight='bold', color='#172139')
    ax.text(3.7, 3.5, "* Navegacion por menu lateral", fontsize=8, color='#334155')
    ax.text(3.7, 3.2, "* Proceso Administrativo", fontsize=8, color='#334155')
    ax.text(3.7, 2.9, "* Proceso Operativo", fontsize=8, color='#334155')
    ax.text(3.7, 2.6, "* Escala de calificacion: 1 a 5", fontsize=8, color='#334155')
    ax.text(3.7, 2.3, "* Resultados con graficos", fontsize=8, color='#334155')
    ax.text(3.7, 2.0, "* Guardar y cargar evaluaciones", fontsize=8, color='#334155')

    card2 = FancyBboxPatch((8.5, 1.5), 3.0, 2.7, boxstyle="round,pad=0.1",
                            facecolor='white', edgecolor='#E2E8F0')
    ax.add_patch(card2)
    border = plt.Rectangle((8.5, 4.15), 3.0, 0.05, color='#667EEA')
    ax.add_patch(border)
    ax.text(10.0, 3.9, "Niveles de Prioridad", fontsize=10, fontweight='bold',
            color='#172139', ha='center')

    for y, color, text in [(3.4, '#10B981', 'BAJA >= 75%'),
                            (2.9, '#F59E0B', 'MEDIA 60-74%'),
                            (2.4, '#EF4444', 'ALTA < 60%')]:
        rect = FancyBboxPatch((8.7, y - 0.15), 2.6, 0.35, boxstyle="round,pad=0.03",
                               facecolor=color, edgecolor='none', alpha=0.15)
        ax.add_patch(rect)
        ax.plot([8.7, 8.7], [y - 0.15, y + 0.2], color=color, linewidth=3)
        ax.text(9.0, y + 0.02, text, fontsize=8, color='#172139', va='center')

    card3 = FancyBboxPatch((8.5, 0.3), 3.0, 1.0, boxstyle="round,pad=0.1",
                            facecolor='white', edgecolor='#E2E8F0')
    ax.add_patch(card3)
    ax.text(10.0, 1.05, "Progreso", fontsize=10, fontweight='bold',
            color='#172139', ha='center')
    ax.text(10.0, 0.7, "0.0%", fontsize=18, fontweight='bold',
            color='#667EEA', ha='center')

    ax.text(7.45, 6.75, "Herramienta de Evaluacion - Gestion Por Procesos (GPP)",
            fontsize=13, fontweight='bold', color='#172139', ha='center', va='center')

    plt.tight_layout()
    plt.savefig(os.path.join(IMG_DIR, "01_pantalla_inicio.png"), dpi=150, bbox_inches='tight',
                facecolor='#F8FAFC')
    plt.close()
    print("  [OK] Imagen: Pantalla de Inicio")


def crear_imagen_sidebar():
    """Crea imagen del menú lateral de navegación"""
    fig, ax = plt.subplots(1, 1, figsize=(4, 8))
    ax.set_xlim(0, 4)
    ax.set_ylim(0, 8)
    ax.axis('off')

    # Fondo sidebar con gradiente
    for i in range(40):
        y = i * (8 / 40)
        r = 0.09 + (i / 40) * 0.05
        g = 0.13 + (i / 40) * 0.2
        b = 0.22 + (i / 40) * 0.12
        rect = plt.Rectangle((0, y), 4, 8/40, color=(r, g, b))
        ax.add_patch(rect)

    ax.text(2, 7.5, "Navegacion", fontsize=16, fontweight='bold',
            color='white', ha='center', va='center')

    # Línea separadora
    ax.plot([0.3, 3.7], [7.1, 7.1], color='white', alpha=0.3, linewidth=1)

    menu_items = [
        ("[1]", "Inicio", True),
        ("[2]", "Proceso Administrativo", False),
        ("[3]", "Proceso Operativo", False),
        ("[4]", "Resultados Generales", False),
        ("[5]", "Prioridades", False),
        ("[6]", "Guardar/Cargar", False),
    ]

    for i, (icon, label, selected) in enumerate(menu_items):
        y = 6.5 - i * 0.9
        if selected:
            rect = FancyBboxPatch((0.3, y - 0.3), 3.4, 0.6, boxstyle="round,pad=0.05",
                                   facecolor='#00AEAC', edgecolor=(0, 0.68, 0.67, 0.4))
        else:
            rect = FancyBboxPatch((0.3, y - 0.3), 3.4, 0.6, boxstyle="round,pad=0.05",
                                   facecolor='white', edgecolor='none', alpha=0.08)
        ax.add_patch(rect)
        ax.text(0.7, y, icon, fontsize=10, color='#00AEAC' if selected else 'white',
                ha='left', va='center', fontweight='bold')
        ax.text(1.3, y, label, fontsize=11, color='white', ha='left', va='center',
                fontweight='bold' if selected else 'normal')

    # Nota al pie
    ax.text(2, 1.5, "Seleccione una\nopcion del menu\npara navegar", fontsize=9,
            color='white', alpha=0.7, ha='center', va='center', style='italic')

    plt.tight_layout()
    plt.savefig(os.path.join(IMG_DIR, "02_menu_navegacion.png"), dpi=150, bbox_inches='tight')
    plt.close()
    print("  [OK] Imagen: Menu de Navegacion")


def crear_imagen_proceso_administrativo():
    """Crea imagen ilustrativa de la evaluación del proceso administrativo"""
    fig, ax = plt.subplots(1, 1, figsize=(12, 8))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis('off')
    fig.patch.set_facecolor('#F8FAFC')

    # Título
    ax.text(6, 7.6, "Evaluacion del Proceso Administrativo", fontsize=16,
            fontweight='bold', color='#172139', ha='center')

    # Info bar
    info_bar = FancyBboxPatch((0.5, 7.0), 11.0, 0.4, boxstyle="round,pad=0.05",
                               facecolor='#E0F2FE', edgecolor='#06B6D4', alpha=0.8)
    ax.add_patch(info_bar)
    ax.text(6, 7.2, "ℹ️ Evalúe cada variable en escala de 1 a 5 (1=No cumple, 5=Cumple plenamente)",
            fontsize=9, color='#0369A1', ha='center', va='center')

    # Sección PLANEACIÓN
    section1 = FancyBboxPatch((0.5, 4.2), 11.0, 2.5, boxstyle="round,pad=0.1",
                               facecolor='white', edgecolor='#E2E8F0')
    ax.add_patch(section1)
    ax.text(1.0, 6.4, "PLANEACIÓN", fontsize=13, fontweight='bold', color='#172139')

    # Expander simulado
    exp1 = FancyBboxPatch((0.7, 5.3), 10.6, 0.8, boxstyle="round,pad=0.05",
                           facecolor='#F8FAFC', edgecolor='#E2E8F0')
    ax.add_patch(exp1)
    ax.text(1.0, 5.9, " Análisis del contexto", fontsize=10, fontweight='bold', color='#172139')

    # Pregunta con selectbox
    ax.text(1.2, 5.5, "¿La empresa realiza un análisis externo como base para su planeación?",
            fontsize=8, color='#334155')

    # Selectbox simulado
    sel1 = FancyBboxPatch((9.5, 5.35), 1.4, 0.3, boxstyle="round,pad=0.03",
                           facecolor='white', edgecolor='#CBD5E1')
    ax.add_patch(sel1)
    ax.text(10.2, 5.5, "▼  3", fontsize=9, color='#172139', ha='center', va='center')

    # Expander 2
    exp2 = FancyBboxPatch((0.7, 4.4), 10.6, 0.7, boxstyle="round,pad=0.05",
                           facecolor='#F8FAFC', edgecolor='#E2E8F0')
    ax.add_patch(exp2)
    ax.text(1.0, 4.95, " Existencia de un plan estratégico", fontsize=10,
            fontweight='bold', color='#172139')
    ax.text(1.2, 4.6, "¿La misión de la empresa está claramente definida y comunicada?",
            fontsize=8, color='#334155')
    sel2 = FancyBboxPatch((9.5, 4.45), 1.4, 0.3, boxstyle="round,pad=0.03",
                           facecolor='white', edgecolor='#CBD5E1')
    ax.add_patch(sel2)
    ax.text(10.2, 4.6, "▼  4", fontsize=9, color='#172139', ha='center', va='center')

    # Sección ORGANIZACIÓN
    section2 = FancyBboxPatch((0.5, 1.5), 11.0, 2.4, boxstyle="round,pad=0.1",
                               facecolor='white', edgecolor='#E2E8F0')
    ax.add_patch(section2)
    ax.text(1.0, 3.6, "ORGANIZACIÓN", fontsize=13, fontweight='bold', color='#172139')

    exp3 = FancyBboxPatch((0.7, 2.8), 10.6, 0.6, boxstyle="round,pad=0.05",
                           facecolor='#F8FAFC', edgecolor='#E2E8F0')
    ax.add_patch(exp3)
    ax.text(1.0, 3.25, " Estructura organizativa", fontsize=10,
            fontweight='bold', color='#172139')
    ax.text(1.2, 2.95, "¿El establecimiento cuenta con un organigrama claramente definido?",
            fontsize=8, color='#334155')
    sel3 = FancyBboxPatch((9.5, 2.82), 1.4, 0.3, boxstyle="round,pad=0.03",
                           facecolor='white', edgecolor='#CBD5E1')
    ax.add_patch(sel3)
    ax.text(10.2, 2.97, "▼  5", fontsize=9, color='#172139', ha='center', va='center')

    exp4 = FancyBboxPatch((0.7, 1.7), 10.6, 0.8, boxstyle="round,pad=0.05",
                           facecolor='#F8FAFC', edgecolor='#E2E8F0')
    ax.add_patch(exp4)
    ax.text(1.0, 2.25, " Procesos documentados", fontsize=10, fontweight='bold', color='#172139')
    ax.text(1.2, 1.95, "¿Existen perfiles de cargo definidos para cada función?",
            fontsize=8, color='#334155')
    sel4 = FancyBboxPatch((9.5, 1.82), 1.4, 0.3, boxstyle="round,pad=0.03",
                           facecolor='white', edgecolor='#CBD5E1')
    ax.add_patch(sel4)
    ax.text(10.2, 1.97, "▼  2", fontsize=9, color='#172139', ha='center', va='center')

    # Botón guardar
    btn = FancyBboxPatch((3.5, 0.3), 5.0, 0.7, boxstyle="round,pad=0.1",
                          facecolor='#172139', edgecolor='none')
    ax.add_patch(btn)
    ax.text(6, 0.65, " Guardar Evaluación del Proceso Administrativo",
            fontsize=11, fontweight='bold', color='white', ha='center', va='center')

    plt.tight_layout()
    plt.savefig(os.path.join(IMG_DIR, "03_proceso_administrativo.png"), dpi=150,
                bbox_inches='tight', facecolor='#F8FAFC')
    plt.close()
    print("  [OK] Imagen: Proceso Administrativo")


def crear_imagen_proceso_operativo():
    """Crea imagen del proceso operativo"""
    fig, ax = plt.subplots(1, 1, figsize=(12, 7))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis('off')
    fig.patch.set_facecolor('#F8FAFC')

    ax.text(6, 6.6, "️ Evaluación del Proceso Operativo", fontsize=16,
            fontweight='bold', color='#172139', ha='center')

    info = FancyBboxPatch((0.5, 6.0), 11.0, 0.4, boxstyle="round,pad=0.05",
                           facecolor='#E0F2FE', edgecolor='#06B6D4', alpha=0.8)
    ax.add_patch(info)
    ax.text(6, 6.2, "ℹ️ Evalúe cada variable de 1 a 5",
            fontsize=9, color='#0369A1', ha='center')

    # LOGÍSTICA DE COMPRAS
    s1 = FancyBboxPatch((0.5, 3.8), 11.0, 2.0, boxstyle="round,pad=0.1",
                          facecolor='white', edgecolor='#E2E8F0')
    ax.add_patch(s1)
    ax.text(1.0, 5.5, "LOGÍSTICA DE COMPRAS", fontsize=13, fontweight='bold', color='#172139')

    for j, (txt, val) in enumerate([
        ("¿La planeación de compras se realiza de manera organizada?", "4"),
        ("¿Se utilizan herramientas de control para la gestión?", "3"),
        ("¿Se cuenta con condiciones adecuadas de almacenamiento?", "5")
    ]):
        y = 5.1 - j * 0.4
        ax.text(1.2, y, txt, fontsize=8, color='#334155')
        sel = FancyBboxPatch((9.5, y - 0.15), 1.4, 0.3, boxstyle="round,pad=0.03",
                              facecolor='white', edgecolor='#CBD5E1')
        ax.add_patch(sel)
        ax.text(10.2, y, f"▼  {val}", fontsize=9, color='#172139', ha='center', va='center')

    # GESTIÓN DE PRODUCCIÓN
    s2 = FancyBboxPatch((0.5, 1.0), 11.0, 2.5, boxstyle="round,pad=0.1",
                          facecolor='white', edgecolor='#E2E8F0')
    ax.add_patch(s2)
    ax.text(1.0, 3.2, "GESTIÓN DE PRODUCCIÓN", fontsize=13, fontweight='bold', color='#172139')

    for j, (txt, val) in enumerate([
        ("¿Los equipos e instalaciones son adecuados para la producción?", "3"),
        ("¿Se realizan actividades de mantenimiento en equipos?", "4"),
        ("¿El diseño de la cocina favorece un trabajo ordenado?", "2"),
        ("¿Se planea adecuadamente los niveles de producción?", "3")
    ]):
        y = 2.8 - j * 0.4
        ax.text(1.2, y, txt, fontsize=8, color='#334155')
        sel = FancyBboxPatch((9.5, y - 0.15), 1.4, 0.3, boxstyle="round,pad=0.03",
                              facecolor='white', edgecolor='#CBD5E1')
        ax.add_patch(sel)
        ax.text(10.2, y, f"▼  {val}", fontsize=9, color='#172139', ha='center', va='center')

    # Botón
    btn = FancyBboxPatch((3.5, 0.2), 5.0, 0.6, boxstyle="round,pad=0.1",
                          facecolor='#172139', edgecolor='none')
    ax.add_patch(btn)
    ax.text(6, 0.5, " Guardar Evaluación del Proceso Operativo",
            fontsize=11, fontweight='bold', color='white', ha='center', va='center')

    plt.tight_layout()
    plt.savefig(os.path.join(IMG_DIR, "04_proceso_operativo.png"), dpi=150,
                bbox_inches='tight', facecolor='#F8FAFC')
    plt.close()
    print("  [OK] Imagen: Proceso Operativo")


def crear_imagen_resultados():
    """Crea imagen de resultados generales con gráficos"""
    fig = plt.figure(figsize=(12, 9))
    fig.patch.set_facecolor('#F8FAFC')

    # Título
    fig.text(0.5, 0.96, " Resultados Generales de la Evaluación", fontsize=16,
             fontweight='bold', color='#172139', ha='center')

    # --- Gauge chart simulado ---
    ax1 = fig.add_axes([0.05, 0.55, 0.4, 0.38])
    theta = np.linspace(np.pi, 0, 100)
    # Fondo del gauge
    for i in range(len(theta)-1):
        pct = (i / len(theta)) * 100
        if pct < 60:
            c = '#E9901E'
        elif pct < 75:
            c = '#FBBB28'
        else:
            c = '#B4C42C'
        ax1.plot([np.cos(theta[i]), np.cos(theta[i+1])],
                 [np.sin(theta[i]), np.sin(theta[i+1])],
                 color=c, linewidth=20, solid_capstyle='round', alpha=0.3)

    # Valor actual (68.5%)
    val = 68.5
    val_theta = np.pi * (1 - val/100)
    # Aguja
    ax1.annotate('', xy=(np.cos(val_theta)*0.7, np.sin(val_theta)*0.7),
                 xytext=(0, 0),
                 arrowprops=dict(arrowstyle='->', color='#172139', lw=3))
    ax1.text(0, -0.2, f"{val}%", fontsize=28, fontweight='bold',
             color='#172139', ha='center', va='center')
    ax1.text(0, -0.45, "Cumplimiento General", fontsize=10, color='#64748B', ha='center')
    ax1.set_xlim(-1.3, 1.3)
    ax1.set_ylim(-0.6, 1.3)
    ax1.axis('off')
    ax1.set_title(" Nivel de Cumplimiento", fontsize=11, fontweight='bold',
                   color='#172139', pad=10)

    # --- Métricas ---
    ax_m = fig.add_axes([0.5, 0.72, 0.47, 0.2])
    ax_m.axis('off')
    metrics = [("Cumplimiento\nGeneral", "68.5%", "#667EEA"),
               ("Calificación\nPromedio", "3.43/5.0", "#10B981"),
               ("Nivel de\nPrioridad", " MEDIA", "#F59E0B")]
    for i, (label, value, color) in enumerate(metrics):
        x = 0.17 + i * 0.33
        rect = FancyBboxPatch((x - 0.13, 0.05), 0.26, 0.9, boxstyle="round,pad=0.05",
                               facecolor='white', edgecolor=color, linewidth=2,
                               transform=ax_m.transAxes)
        ax_m.add_patch(rect)
        ax_m.text(x, 0.7, value, fontsize=16, fontweight='bold', color=color,
                  ha='center', va='center', transform=ax_m.transAxes)
        ax_m.text(x, 0.3, label, fontsize=8, color='#64748B',
                  ha='center', va='center', transform=ax_m.transAxes)

    # --- Gráfico de barras horizontales ---
    ax2 = fig.add_axes([0.55, 0.08, 0.42, 0.58])
    areas = ['PA: Planeación', 'PA: Organización', 'PA: Dirección', 'PA: Control',
             'PO: Log. Compras', 'PO: Gest. Producción', 'PO: Log. Externa']
    valores = [72, 65, 80, 55, 70, 62, 78]
    colores = ['#FBBB28', '#FBBB28', '#B4C42C', '#E9901E', '#FBBB28', '#FBBB28', '#B4C42C']

    bars = ax2.barh(areas, valores, color=colores, edgecolor='white', linewidth=1.5, height=0.6)
    ax2.set_xlim(0, 105)
    ax2.set_xlabel("Cumplimiento (%)", fontsize=9, color='#172139')
    ax2.set_title(" Cumplimiento por Área", fontsize=12, fontweight='bold',
                   color='#172139', pad=10)

    # Líneas de referencia
    ax2.axvline(x=75, color='#B4C42C', linestyle='--', alpha=0.5, linewidth=1)
    ax2.axvline(x=60, color='#E9901E', linestyle='--', alpha=0.5, linewidth=1)

    for bar, val in zip(bars, valores):
        ax2.text(val + 1, bar.get_y() + bar.get_height()/2, f"{val}%",
                 fontsize=9, fontweight='bold', va='center', color='#172139')

    ax2.tick_params(axis='y', labelsize=8)
    ax2.tick_params(axis='x', labelsize=8)
    ax2.spines['top'].set_visible(False)
    ax2.spines['right'].set_visible(False)
    ax2.set_facecolor('#FAFBFC')

    # --- Gráfico radar ---
    ax3 = fig.add_axes([0.02, 0.02, 0.45, 0.5], polar=True)
    categories = ['Planeación', 'Organización', 'Dirección', 'Control',
                  'Log. Compras', 'Gest. Producción', 'Log. Externa']
    values = [72, 65, 80, 55, 70, 62, 78]
    values_closed = values + [values[0]]

    angles = np.linspace(0, 2 * np.pi, len(categories), endpoint=False).tolist()
    angles += angles[:1]

    # Área ideal
    ax3.fill(angles, [100]*len(angles), color='#B4C42C', alpha=0.05)
    ax3.plot(angles, [100]*len(angles), color='#B4C42C', linewidth=1, linestyle=':', alpha=0.4)

    # Meta mínima
    ax3.fill(angles, [75]*len(angles), color='#FBBB28', alpha=0.05)
    ax3.plot(angles, [75]*len(angles), color='#FBBB28', linewidth=1, linestyle='--', alpha=0.5)

    # Valores actuales
    ax3.fill(angles, values_closed, color='#172139', alpha=0.25)
    ax3.plot(angles, values_closed, color='#172139', linewidth=2.5, marker='o', markersize=6)

    ax3.set_xticks(angles[:-1])
    ax3.set_xticklabels(categories, fontsize=7, color='#172139')
    ax3.set_ylim(0, 100)
    ax3.set_title(" Análisis Radial", fontsize=11, fontweight='bold',
                   color='#172139', pad=20)
    ax3.tick_params(axis='y', labelsize=7)

    plt.savefig(os.path.join(IMG_DIR, "05_resultados_generales.png"), dpi=150,
                bbox_inches='tight', facecolor='#F8FAFC')
    plt.close()
    print("  [OK] Imagen: Resultados Generales")


def crear_imagen_prioridades():
    """Crea imagen de análisis de prioridades"""
    fig = plt.figure(figsize=(12, 7))
    fig.patch.set_facecolor('#F8FAFC')

    fig.text(0.5, 0.95, " Análisis de Prioridades", fontsize=16,
             fontweight='bold', color='#172139', ha='center')

    # Gráfico de dona
    ax1 = fig.add_axes([0.05, 0.15, 0.45, 0.7])
    sizes = [4, 5, 7]
    colors = ['#EF4444', '#F59E0B', '#10B981']
    labels = ['Alta (4)', 'Media (5)', 'Baja (7)']
    explode = (0.05, 0, 0)

    wedges, texts, autotexts = ax1.pie(sizes, explode=explode, labels=labels,
                                        colors=colors, autopct='%1.0f%%',
                                        startangle=90, pctdistance=0.8,
                                        wedgeprops=dict(width=0.5, edgecolor='white', linewidth=3))

    for text in texts:
        text.set_fontsize(10)
        text.set_fontweight('bold')
    for autotext in autotexts:
        autotext.set_fontsize(10)
        autotext.set_fontweight('bold')
        autotext.set_color('white')

    ax1.text(0, 0, "16\nElementos\nTotales", ha='center', va='center',
             fontsize=14, fontweight='bold', color='#172139')
    ax1.set_title(" Distribución de Prioridades", fontsize=12, fontweight='bold',
                   color='#172139', pad=15)

    # Cards de conteo
    ax2 = fig.add_axes([0.55, 0.15, 0.42, 0.7])
    ax2.axis('off')

    card_info = [
        (" Prioridad Alta", "4", "#EF4444", "Acción inmediata requerida\n(<60% cumplimiento)"),
        (" Prioridad Media", "5", "#F59E0B", "Atender en mediano plazo\n(60-74% cumplimiento)"),
        (" Prioridad Baja", "7", "#10B981", "No requiere acción inmediata\n(≥75% cumplimiento)")
    ]

    for i, (label, count, color, desc) in enumerate(card_info):
        y = 0.82 - i * 0.32
        rect = FancyBboxPatch((0.02, y - 0.1), 0.96, 0.28, boxstyle="round,pad=0.03",
                               facecolor='white', edgecolor=color, linewidth=2,
                               transform=ax2.transAxes)
        ax2.add_patch(rect)

        # Borde superior coloreado
        border = plt.Rectangle((0.02, y + 0.16), 0.96, 0.02, color=color,
                                transform=ax2.transAxes)
        ax2.add_patch(border)

        ax2.text(0.15, y + 0.1, label, fontsize=10, fontweight='bold',
                 color='#172139', va='center', transform=ax2.transAxes)
        ax2.text(0.85, y + 0.06, count, fontsize=26, fontweight='bold',
                 color=color, ha='center', va='center', transform=ax2.transAxes)
        ax2.text(0.15, y - 0.03, desc, fontsize=7, color='#64748B',
                 va='center', transform=ax2.transAxes)

    plt.savefig(os.path.join(IMG_DIR, "06_prioridades.png"), dpi=150,
                bbox_inches='tight', facecolor='#F8FAFC')
    plt.close()
    print("  [OK] Imagen: Prioridades")


def crear_imagen_guardar_cargar():
    """Crea imagen de la sección guardar/cargar"""
    fig, ax = plt.subplots(1, 1, figsize=(12, 6))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6)
    ax.axis('off')
    fig.patch.set_facecolor('#F8FAFC')

    ax.text(6, 5.6, " Guardar y Cargar Evaluaciones", fontsize=16,
            fontweight='bold', color='#172139', ha='center')

    # Card Guardar
    c1 = FancyBboxPatch((0.3, 2.2), 5.2, 3.0, boxstyle="round,pad=0.1",
                          facecolor='white', edgecolor='#E2E8F0')
    ax.add_patch(c1)
    border1 = plt.Rectangle((0.3, 5.15), 5.2, 0.05, color='#10B981')
    ax.add_patch(border1)

    ax.text(2.9, 4.9, " Guardar Evaluación Actual", fontsize=12,
            fontweight='bold', color='#172139', ha='center')
    ax.text(2.9, 4.4, "Guarda la evaluación actual en formato\nJSON para consulta posterior.",
            fontsize=9, color='#64748B', ha='center')

    btn1 = FancyBboxPatch((1.0, 3.3), 3.8, 0.6, boxstyle="round,pad=0.1",
                           facecolor='#172139', edgecolor='none')
    ax.add_patch(btn1)
    ax.text(2.9, 3.6, "Guardar Evaluación", fontsize=11, fontweight='bold',
            color='white', ha='center', va='center')

    # Mensaje de éxito
    success = FancyBboxPatch((0.7, 2.4), 4.4, 0.6, boxstyle="round,pad=0.05",
                              facecolor='#F0FDF4', edgecolor='#10B981')
    ax.add_patch(success)
    ax.text(2.9, 2.7, " Evaluación guardada exitosamente", fontsize=9,
            color='#166534', ha='center', va='center')

    # Card Cargar
    c2 = FancyBboxPatch((6.5, 2.2), 5.2, 3.0, boxstyle="round,pad=0.1",
                          facecolor='white', edgecolor='#E2E8F0')
    ax.add_patch(c2)
    border2 = plt.Rectangle((6.5, 5.15), 5.2, 0.05, color='#667EEA')
    ax.add_patch(border2)

    ax.text(9.1, 4.9, " Cargar Evaluación Anterior", fontsize=12,
            fontweight='bold', color='#172139', ha='center')

    # Select box simulado
    sel = FancyBboxPatch((7.0, 4.1), 4.2, 0.5, boxstyle="round,pad=0.05",
                          facecolor='white', edgecolor='#CBD5E1')
    ax.add_patch(sel)
    ax.text(9.1, 4.35, "evaluacion_20251017_152909    ▼",
            fontsize=9, color='#172139', ha='center', va='center')

    btn2 = FancyBboxPatch((7.2, 3.3), 3.8, 0.6, boxstyle="round,pad=0.1",
                           facecolor='#172139', edgecolor='none')
    ax.add_patch(btn2)
    ax.text(9.1, 3.6, "Cargar Evaluación", fontsize=11, fontweight='bold',
            color='white', ha='center', va='center')

    # Sección Exportar
    c3 = FancyBboxPatch((0.3, 0.3), 11.4, 1.6, boxstyle="round,pad=0.1",
                          facecolor='white', edgecolor='#E2E8F0')
    ax.add_patch(c3)
    border3 = plt.Rectangle((0.3, 1.85), 11.4, 0.05, color='#F59E0B')
    ax.add_patch(border3)

    ax.text(6, 1.6, " Exportar Resultados a Excel", fontsize=12,
            fontweight='bold', color='#172139', ha='center')
    ax.text(6, 1.2, "Exporte los resultados de su evaluación a un archivo Excel (.xlsx) para reportes formales.",
            fontsize=9, color='#64748B', ha='center')

    btn3 = FancyBboxPatch((4.0, 0.4), 4.0, 0.5, boxstyle="round,pad=0.1",
                           facecolor='#172139', edgecolor='none')
    ax.add_patch(btn3)
    ax.text(6, 0.65, "Exportar a Excel", fontsize=11, fontweight='bold',
            color='white', ha='center', va='center')

    plt.tight_layout()
    plt.savefig(os.path.join(IMG_DIR, "07_guardar_cargar.png"), dpi=150,
                bbox_inches='tight', facecolor='#F8FAFC')
    plt.close()
    print("  [OK] Imagen: Guardar/Cargar")


def crear_imagen_escala_calificacion():
    """Crea imagen de la escala de calificación"""
    fig, ax = plt.subplots(1, 1, figsize=(10, 4))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4)
    ax.axis('off')
    fig.patch.set_facecolor('#F8FAFC')

    ax.text(5, 3.7, " Escala de Calificación (1 a 5)", fontsize=14,
            fontweight='bold', color='#172139', ha='center')

    escala = [
        (1, "No cumple", "#EF4444", "El establecimiento no cumple\ncon el criterio evaluado"),
        (2, "Cumple parcialmente", "#F97316", "Existe un cumplimiento\nmínimo o incipiente"),
        (3, "Cumple aceptablemente", "#F59E0B", "Cumplimiento aceptable\ncon mejoras posibles"),
        (4, "Cumple en gran medida", "#84CC16", "Cumplimiento alto con\noportunidades menores"),
        (5, "Cumple plenamente", "#10B981", "El criterio se cumple\nde forma completa")
    ]

    for i, (num, label, color, desc) in enumerate(escala):
        x = 0.3 + i * 1.95
        # Card
        card = FancyBboxPatch((x, 0.3), 1.7, 3.0, boxstyle="round,pad=0.08",
                               facecolor='white', edgecolor=color, linewidth=2)
        ax.add_patch(card)
        # Borde superior
        top_border = plt.Rectangle((x, 3.2), 1.7, 0.1, color=color)
        ax.add_patch(top_border)
        # Número
        circle = plt.Circle((x + 0.85, 2.7), 0.3, color=color, alpha=0.15)
        ax.add_patch(circle)
        ax.text(x + 0.85, 2.7, str(num), fontsize=20, fontweight='bold',
                color=color, ha='center', va='center')
        # Label
        ax.text(x + 0.85, 2.1, label, fontsize=8, fontweight='bold',
                color='#172139', ha='center', va='center')
        # Descripción
        ax.text(x + 0.85, 1.3, desc, fontsize=6.5, color='#64748B',
                ha='center', va='center')

    plt.tight_layout()
    plt.savefig(os.path.join(IMG_DIR, "08_escala_calificacion.png"), dpi=150,
                bbox_inches='tight', facecolor='#F8FAFC')
    plt.close()
    print("  [OK] Imagen: Escala de Calificación")


def crear_imagen_instalacion():
    """Crea imagen del proceso de instalación"""
    fig, ax = plt.subplots(1, 1, figsize=(10, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis('off')
    fig.patch.set_facecolor('#1E1E2E')

    # Terminal simulada
    terminal = FancyBboxPatch((0.3, 0.3), 9.4, 4.4, boxstyle="round,pad=0.1",
                               facecolor='#0D1117', edgecolor='#30363D', linewidth=2)
    ax.add_patch(terminal)

    # Barra de título de terminal
    title_bar = FancyBboxPatch((0.3, 4.2), 9.4, 0.5, boxstyle="round,pad=0.05",
                                facecolor='#161B22', edgecolor='none')
    ax.add_patch(title_bar)

    # Botones de ventana
    for i, c in enumerate(['#FF5F57', '#FEBC2E', '#28C840']):
        circle = plt.Circle((0.7 + i * 0.35, 4.45), 0.08, color=c)
        ax.add_patch(circle)

    ax.text(5, 4.45, "PowerShell - Instalador GPP", fontsize=9, color='#8B949E',
            ha='center', va='center')

    # Contenido de la terminal
    lines = [
        ("", "========================================", "#58A6FF"),
        ("", "Instalador de la Aplicacion GPP", "#58A6FF"),
        ("", "========================================", "#58A6FF"),
        ("", "", ""),
        ("", "Verificando Python...", "#E6EDF3"),
        ("", "Python 3.12.0", "#3FB950"),
        ("", "", ""),
        ("", "Instalando dependencias...", "#E6EDF3"),
        ("", "[OK] streamlit       - version 1.28.0", "#3FB950"),
        ("", "[OK] pandas          - version 2.1.2", "#3FB950"),
        ("", "[OK] plotly          - version 5.18.0", "#3FB950"),
        ("", "[OK] numpy           - version 1.26.1", "#3FB950"),
        ("", "[OK] openpyxl        - version 3.1.2", "#3FB950"),
        ("", "", ""),
        ("", "========================================", "#58A6FF"),
        ("", "Instalacion completada exitosamente!", "#3FB950"),
        ("", "========================================", "#58A6FF"),
    ]

    for i, (prefix, text, color) in enumerate(lines):
        y = 4.0 - i * 0.22
        if text:
            ax.text(0.6, y, f"{prefix}{text}", fontsize=7.5, color=color,
                    fontfamily='monospace', va='center')

    plt.tight_layout()
    plt.savefig(os.path.join(IMG_DIR, "09_instalacion.png"), dpi=150,
                bbox_inches='tight', facecolor='#1E1E2E')
    plt.close()
    print("  [OK] Imagen: Proceso de Instalación")


def crear_imagen_flujo_trabajo():
    """Crea diagrama de flujo de trabajo"""
    fig, ax = plt.subplots(1, 1, figsize=(11, 5))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 5)
    ax.axis('off')
    fig.patch.set_facecolor('#F8FAFC')

    ax.text(5.5, 4.7, " Flujo de Trabajo de la Aplicación GPP", fontsize=14,
            fontweight='bold', color='#172139', ha='center')

    steps = [
        (1.0, 2.5, "1. Instalar\nDependencias", "#667EEA", "instalar.bat"),
        (3.2, 2.5, "2. Ejecutar\nAplicación", "#06B6D4", "ejecutar.bat"),
        (5.4, 2.5, "3. Evaluar\nProcesos", "#10B981", "PA y PO"),
        (7.6, 2.5, "4. Ver\nResultados", "#F59E0B", "Gráficos"),
        (9.8, 2.5, "5. Guardar\ny Exportar", "#EF4444", "JSON/Excel"),
    ]

    for x, y, label, color, sub in steps:
        # Círculo principal
        circle = plt.Circle((x, y), 0.7, color=color, alpha=0.15)
        ax.add_patch(circle)
        circle_border = plt.Circle((x, y), 0.7, fill=False, edgecolor=color, linewidth=2.5)
        ax.add_patch(circle_border)

        ax.text(x, y + 0.1, label, fontsize=8, fontweight='bold',
                color='#172139', ha='center', va='center')
        ax.text(x, y - 0.55, sub, fontsize=7, color=color,
                ha='center', va='center', style='italic')

    # Flechas de conexión
    for i in range(len(steps) - 1):
        x1 = steps[i][0] + 0.75
        x2 = steps[i + 1][0] - 0.75
        ax.annotate('', xy=(x2, 2.5), xytext=(x1, 2.5),
                     arrowprops=dict(arrowstyle='->', color='#94A3B8', lw=2,
                                      connectionstyle='arc3,rad=0'))

    # Nota inferior
    note = FancyBboxPatch((1.5, 0.3), 8.0, 0.8, boxstyle="round,pad=0.1",
                           facecolor='#E0F2FE', edgecolor='#06B6D4', alpha=0.8)
    ax.add_patch(note)
    ax.text(5.5, 0.7, " El paso 1 solo es necesario la primera vez. Los pasos 3-5 pueden repetirse las veces que necesite.",
            fontsize=8, color='#0369A1', ha='center', va='center')

    plt.tight_layout()
    plt.savefig(os.path.join(IMG_DIR, "10_flujo_trabajo.png"), dpi=150,
                bbox_inches='tight', facecolor='#F8FAFC')
    plt.close()
    print("  [OK] Imagen: Flujo de Trabajo")


def crear_imagen_arquitectura():
    """Crea diagrama de arquitectura de archivos"""
    fig, ax = plt.subplots(1, 1, figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis('off')
    fig.patch.set_facecolor('#F8FAFC')

    ax.text(5, 5.7, " Estructura de Archivos de la Aplicación", fontsize=14,
            fontweight='bold', color='#172139', ha='center')

    # Carpeta principal
    main_folder = FancyBboxPatch((0.5, 0.3), 9.0, 5.0, boxstyle="round,pad=0.1",
                                  facecolor='white', edgecolor='#E2E8F0', linewidth=1.5)
    ax.add_patch(main_folder)

    ax.text(5, 5.15, " Aplicacion_GPP/", fontsize=13, fontweight='bold',
            color='#172139', ha='center')

    files = [
        ("", "app.py", "Aplicación principal (Python/Streamlit)", "#667EEA"),
        ("", "ejecutar.bat", "Script para iniciar la aplicación", "#10B981"),
        ("", "instalar.bat", "Script de instalación de dependencias", "#06B6D4"),
        ("", "verificar_instalacion.py", "Verifica que todo esté instalado", "#F59E0B"),
        ("", "requirements.txt", "Lista de dependencias necesarias", "#8B5CF6"),
        ("", "LEEME_PRIMERO.txt", "Información y guía inicial", "#EC4899"),
        ("", "INSTRUCCIONES_RAPIDAS.txt", "Guía rápida de uso", "#EF4444"),
        ("", "COMO_ABRIR_HTML.txt", "Instrucciones versión HTML", "#F97316"),
        ("", "evaluacion_*.json", "Evaluaciones guardadas", "#94A3B8"),
        ("", "resultados_*.xlsx", "Exportaciones a Excel", "#94A3B8"),
    ]

    for i, (icon, name, desc, color) in enumerate(files):
        y = 4.65 - i * 0.44
        # Fondo de fila
        if i % 2 == 0:
            row_bg = plt.Rectangle((0.7, y - 0.15), 8.6, 0.4, color='#F8FAFC')
            ax.add_patch(row_bg)

        # Indicador de color
        indicator = plt.Rectangle((0.8, y - 0.1), 0.06, 0.3, color=color, alpha=0.8)
        ax.add_patch(indicator)

        ax.text(1.1, y + 0.05, f"{icon} {name}", fontsize=9, fontweight='bold',
                color='#172139', va='center')
        ax.text(5.5, y + 0.05, desc, fontsize=8, color='#64748B',
                va='center', style='italic')

    plt.tight_layout()
    plt.savefig(os.path.join(IMG_DIR, "11_estructura_archivos.png"), dpi=150,
                bbox_inches='tight', facecolor='#F8FAFC')
    plt.close()
    print("  [OK] Imagen: Estructura de Archivos")


def generar_todas_las_imagenes():
    """Genera todas las imágenes ilustrativas"""
    print("\n Generando imágenes ilustrativas del programa...")
    crear_imagen_bienvenida()
    crear_imagen_sidebar()
    crear_imagen_proceso_administrativo()
    crear_imagen_proceso_operativo()
    crear_imagen_resultados()
    crear_imagen_prioridades()
    crear_imagen_guardar_cargar()
    crear_imagen_escala_calificacion()
    crear_imagen_instalacion()
    crear_imagen_flujo_trabajo()
    crear_imagen_arquitectura()
    print("   Todas las imágenes generadas exitosamente\n")


# ============================================================
# PARTE 2: GENERACIÓN DEL DOCUMENTO WORD
# ============================================================

def set_cell_shading(cell, color):
    """Aplica color de fondo a una celda de tabla"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False,
                             font_size=11, font_name='Calibri', color=None,
                             alignment=None, space_before=0, space_after=6,
                             first_line_indent=None):
    """Agrega un párrafo formateado al documento"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)
    return p


def add_heading_styled(doc, text, level=1):
    """Agrega un heading con estilo"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.color.rgb = RGBColor(23, 33, 57)
        elif level == 2:
            run.font.color.rgb = RGBColor(25, 88, 85)
        else:
            run.font.color.rgb = RGBColor(16, 129, 129)
    return h


def add_image_with_caption(doc, image_path, caption, width=Inches(6.0)):
    """Agrega una imagen con su caption al documento"""
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=width)

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = cap.add_run(caption)
        run_cap.italic = True
        run_cap.font.size = Pt(9)
        run_cap.font.color.rgb = RGBColor(100, 116, 139)
        cap.paragraph_format.space_after = Pt(12)
    else:
        doc.add_paragraph(f"[Imagen no disponible: {caption}]")


def crear_documento_word():
    """Crea el documento Word completo con el manual de usuario"""
    print("📝 Generando documento Word...")

    doc = Document()

    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ================================================================
    # PORTADA
    # ================================================================
    for _ in range(4):
        doc.add_paragraph()

    # Título principal
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MANUAL DE USUARIO")
    run.bold = True
    run.font.size = Pt(32)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(23, 33, 57)

    # Subtítulo
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Herramienta de Evaluación de\nGestión Por Procesos (GPP)")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(25, 88, 85)

    doc.add_paragraph()

    # Línea decorativa
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("━" * 50)
    run.font.color.rgb = RGBColor(0, 174, 172)
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Subtítulo 2
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Aplicación Web de Evaluación Profesional\nSistema Interactivo con Gráficos y Análisis de Prioridades")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 116, 139)
    run.font.name = 'Calibri'

    doc.add_paragraph()
    doc.add_paragraph()

    # Información de versión
    info_items = [
        ("Versión:", "Premium 2025"),
        ("Tecnología:", "Python + Streamlit + Plotly"),
        ("Fecha:", datetime.now().strftime("%d de %B de %Y")),
        ("Institución:", "CEITTO - Colmayor"),
        ("Proyecto:", "Gastrocolma"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p.add_run(f"{label} ")
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_label.font.color.rgb = RGBColor(23, 33, 57)
        run_value = p.add_run(value)
        run_value.font.size = Pt(11)
        run_value.font.color.rgb = RGBColor(100, 116, 139)

    # Salto de página
    doc.add_page_break()

    # ================================================================
    # TABLA DE CONTENIDOS
    # ================================================================
    add_heading_styled(doc, "Tabla de Contenido", level=1)

    toc_items = [
        "1. Introducción",
        "   1.1. Propósito del documento",
        "   1.2. Descripción general del sistema",
        "   1.3. Audiencia objetivo",
        "2. Requisitos del Sistema",
        "   2.1. Requisitos de hardware",
        "   2.2. Requisitos de software",
        "   2.3. Dependencias del proyecto",
        "3. Instalación y Configuración Inicial",
        "   3.1. Instalación automática",
        "   3.2. Instalación manual",
        "   3.3. Verificación de la instalación",
        "4. Inicio de la Aplicación",
        "   4.1. Método 1: Usando archivos .bat",
        "   4.2. Método 2: Usando comandos",
        "   4.3. Acceso desde el navegador",
        "5. Estructura de la Aplicación",
        "   5.1. Arquitectura de archivos",
        "   5.2. Flujo de trabajo general",
        "6. Pantalla de Inicio",
        "   6.1. Bienvenida y guía de uso",
        "   6.2. Escala de calificación",
        "   6.3. Barra de progreso",
        "   6.4. Niveles de prioridad",
        "7. Evaluación del Proceso Administrativo",
        "   7.1. Planeación",
        "   7.2. Organización",
        "   7.3. Dirección",
        "   7.4. Control",
        "   7.5. Cómo calificar las variables",
        "8. Evaluación del Proceso Operativo",
        "   8.1. Logística de Compras",
        "   8.2. Gestión de Producción",
        "   8.3. Logística Externa",
        "9. Resultados Generales",
        "   9.1. Panel de métricas principales",
        "   9.2. Indicador de cumplimiento (Gauge)",
        "   9.3. Gráfico de barras por área",
        "   9.4. Análisis radial multidimensional",
        "   9.5. Tablas de resultados detallados",
        "10. Análisis de Prioridades",
        "    10.1. Clasificación de prioridades",
        "    10.2. Gráfico de distribución",
        "    10.3. Tablas detalladas por nivel",
        "11. Guardar, Cargar y Exportar",
        "    11.1. Guardar evaluaciones (JSON)",
        "    11.2. Cargar evaluaciones anteriores",
        "    11.3. Exportar a Excel",
        "12. Versión HTML Alternativa",
        "13. Solución de Problemas",
        "14. Recomendaciones de Uso",
        "15. Glosario de Términos",
    ]

    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        if not item.startswith("   "):
            run.bold = True
            run.font.color.rgb = RGBColor(23, 33, 57)
        else:
            run.font.color.rgb = RGBColor(100, 116, 139)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(1)

    doc.add_page_break()

    # ================================================================
    # CAPÍTULO 1: INTRODUCCIÓN
    # ================================================================
    add_heading_styled(doc, "1. Introducción", level=1)

    add_heading_styled(doc, "1.1. Propósito del documento", level=2)
    add_formatted_paragraph(doc,
        "El presente manual de usuario tiene como propósito servir de guía completa y detallada para la correcta "
        "utilización de la Herramienta de Evaluación de Gestión Por Procesos (GPP). Este documento está diseñado para "
        "proporcionar al usuario toda la información necesaria para instalar, configurar y operar la aplicación de "
        "manera eficiente, desde el primer contacto con el sistema hasta el dominio completo de todas sus funcionalidades. "
        "A lo largo de estas páginas, el usuario encontrará instrucciones paso a paso, capturas de pantalla ilustrativas, "
        "explicaciones detalladas de cada módulo y recomendaciones de uso que facilitarán la experiencia con la herramienta.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "1.2. Descripción general del sistema", level=2)
    add_formatted_paragraph(doc,
        "La Herramienta de Evaluación GPP es una aplicación web moderna e interactiva desarrollada con tecnologías de "
        "última generación (Python, Streamlit y Plotly) que transforma el proceso tradicional de evaluación basado en "
        "hojas de cálculo Excel en una experiencia digital intuitiva y visualmente atractiva. La aplicación permite "
        "realizar autoevaluaciones completas de la gestión por procesos en establecimientos gastronómicos y de servicios "
        "de alimentación, identificando fortalezas y oportunidades de mejora a través de un análisis sistemático y "
        "estructurado.",
        font_size=11, first_line_indent=1.0)

    add_formatted_paragraph(doc,
        "El sistema replica y mejora toda la funcionalidad del archivo Excel original \"Herramienta de evaluación GPP_0\", "
        "incorporando características avanzadas como gráficos interactivos de alta calidad, cálculos automáticos en "
        "tiempo real, clasificación inteligente de prioridades, persistencia de datos mediante archivos JSON, exportación "
        "a Excel y una interfaz de usuario profesional con efectos visuales modernos como glassmorphism, animaciones "
        "suaves y una paleta de colores corporativa cuidadosamente seleccionada.",
        font_size=11, first_line_indent=1.0)

    add_formatted_paragraph(doc,
        "La aplicación evalúa dos grandes áreas de la gestión organizacional: el Proceso Administrativo, que comprende "
        "los aspectos de Planeación, Organización, Dirección y Control; y el Proceso Operativo, que abarca la Logística "
        "de Compras, la Gestión de Producción y la Logística Externa. Cada área contiene múltiples elementos y variables "
        "que son calificados en una escala de 1 a 5, donde los resultados se presentan mediante indicadores de "
        "cumplimiento porcentual y niveles de prioridad que orientan la toma de decisiones para la mejora continua.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "1.3. Audiencia objetivo", level=2)
    add_formatted_paragraph(doc,
        "Este manual está dirigido a los propietarios, administradores, gerentes y personal de establecimientos "
        "gastronómicos que deseen evaluar y mejorar sus procesos de gestión. También resulta útil para consultores, "
        "auditores y profesionales del sector que requieran una herramienta estandarizada para el diagnóstico "
        "organizacional. No se requieren conocimientos técnicos avanzados para utilizar la aplicación; sin embargo, "
        "para el proceso de instalación inicial es conveniente contar con nociones básicas de manejo de computador "
        "y navegación web.",
        font_size=11, first_line_indent=1.0)

    doc.add_page_break()

    # ================================================================
    # CAPÍTULO 2: REQUISITOS DEL SISTEMA
    # ================================================================
    add_heading_styled(doc, "2. Requisitos del Sistema", level=1)

    add_heading_styled(doc, "2.1. Requisitos de hardware", level=2)
    add_formatted_paragraph(doc,
        "La aplicación ha sido diseñada para funcionar de manera eficiente en equipos de cómputo de gama media. "
        "Los requisitos mínimos de hardware incluyen un procesador de al menos 1 GHz (se recomienda de doble núcleo "
        "o superior), una memoria RAM mínima de 2 GB (se recomiendan 4 GB para un rendimiento óptimo), y al menos "
        "500 MB de espacio libre en disco duro para la instalación de la aplicación y sus dependencias. Adicionalmente, "
        "se requiere una pantalla con resolución mínima de 1024x768 píxeles, aunque se recomienda una resolución de "
        "1920x1080 o superior para aprovechar al máximo la interfaz de usuario y los gráficos interactivos.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "2.2. Requisitos de software", level=2)
    add_formatted_paragraph(doc,
        "Para ejecutar la aplicación en su versión completa (Streamlit), es indispensable contar con los siguientes "
        "componentes de software instalados en el equipo:",
        font_size=11, first_line_indent=1.0)

    # Tabla de requisitos
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Componente", "Versión Mínima", "Observaciones"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, "172139")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ["Sistema Operativo", "Windows 10/11, macOS, Linux", "Compatible con los principales SO"],
        ["Python", "3.8 o superior", "Se recomienda Python 3.10+"],
        ["Navegador Web", "Chrome, Edge, Firefox", "Se recomienda Google Chrome"],
        ["Conexión a Internet", "Para instalación inicial", "No requerida después de instalar"],
    ]

    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    add_heading_styled(doc, "2.3. Dependencias del proyecto", level=2)
    add_formatted_paragraph(doc,
        "La aplicación depende de varias bibliotecas de Python que se instalan automáticamente durante el proceso "
        "de configuración inicial. Estas dependencias incluyen Streamlit (framework para crear la interfaz web), "
        "Pandas (manipulación y análisis de datos), Plotly (generación de gráficos interactivos), NumPy (cálculos "
        "numéricos y estadísticos) y OpenPyXL (lectura y escritura de archivos Excel). Todas estas bibliotecas se "
        "encuentran listadas en el archivo requirements.txt y se instalan mediante el administrador de paquetes pip "
        "de Python.",
        font_size=11, first_line_indent=1.0)

    doc.add_page_break()

    # ================================================================
    # CAPÍTULO 3: INSTALACIÓN
    # ================================================================
    add_heading_styled(doc, "3. Instalación y Configuración Inicial", level=1)

    add_formatted_paragraph(doc,
        "El proceso de instalación de la aplicación GPP es sencillo y puede realizarse de forma automática o manual. "
        "A continuación se describen ambos métodos en detalle. Es importante destacar que este proceso solo debe "
        "realizarse una única vez; una vez completada la instalación, el usuario podrá ejecutar la aplicación "
        "directamente las veces que lo necesite.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "3.1. Instalación automática (Recomendada)", level=2)
    add_formatted_paragraph(doc,
        "El método de instalación automática es el más sencillo y recomendado para la mayoría de los usuarios. "
        "Consiste en ejecutar el archivo instalar.bat que se encuentra en la carpeta raíz de la aplicación. Este "
        "archivo batch realiza automáticamente las siguientes operaciones: verifica que Python esté correctamente "
        "instalado en el sistema, actualiza el gestor de paquetes pip a su versión más reciente, y descarga e instala "
        "todas las dependencias necesarias listadas en el archivo requirements.txt.",
        font_size=11, first_line_indent=1.0)

    add_formatted_paragraph(doc,
        "Para ejecutar la instalación automática, siga estos pasos: primero, navegue hasta la carpeta de la aplicación "
        "(Aplicacion_GPP) en el explorador de archivos de Windows. Luego, localice el archivo llamado instalar.bat y "
        "haga doble clic sobre él. Se abrirá una ventana de consola de comandos que mostrará el progreso de la "
        "instalación. Espere a que el proceso finalice completamente; el sistema le indicará cuando la instalación "
        "haya sido exitosa. Finalmente, cierre la ventana de consola.",
        font_size=11, first_line_indent=1.0)

    add_image_with_caption(doc, os.path.join(IMG_DIR, "09_instalacion.png"),
                           "Figura 1. Proceso de instalación de dependencias en la consola de comandos.")

    add_heading_styled(doc, "3.2. Instalación manual", level=2)
    add_formatted_paragraph(doc,
        "Si por alguna razón el método automático no funciona correctamente, el usuario puede optar por la instalación "
        "manual siguiendo estos pasos: abra una ventana de PowerShell o Símbolo del sistema (CMD) y navegue hasta la "
        "carpeta de la aplicación utilizando el comando cd seguido de la ruta de la carpeta. Una vez ubicado en el "
        "directorio correcto, ejecute el siguiente comando para instalar todas las dependencias necesarias: "
        "\"py -m pip install -r requirements.txt\". Este comando leerá el archivo requirements.txt y descargará e "
        "instalará automáticamente todas las bibliotecas requeridas.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "3.3. Verificación de la instalación", level=2)
    add_formatted_paragraph(doc,
        "Una vez completada la instalación, ya sea por el método automático o manual, es altamente recomendable "
        "verificar que todos los componentes se hayan instalado correctamente. Para ello, ejecute el script de "
        "verificación incluido en la aplicación mediante el comando: \"py verificar_instalacion.py\". Este script "
        "comprobará la presencia y versión de cada módulo requerido (Streamlit, Pandas, Plotly, NumPy y OpenPyXL) "
        "y mostrará el estado de cada uno con la etiqueta [OK] si está correctamente instalado o [ERROR] si falta "
        "algún componente. Si todos los módulos aparecen con [OK], la aplicación está lista para ser utilizada.",
        font_size=11, first_line_indent=1.0)

    doc.add_page_break()

    # ================================================================
    # CAPÍTULO 4: INICIO DE LA APLICACIÓN
    # ================================================================
    add_heading_styled(doc, "4. Inicio de la Aplicación", level=1)

    add_formatted_paragraph(doc,
        "Una vez que la instalación ha sido completada y verificada exitosamente, la aplicación puede iniciarse "
        "en cualquier momento mediante dos métodos principales. La aplicación se ejecuta como un servidor web local "
        "y se accede a ella a través del navegador de Internet del equipo.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "4.1. Método 1: Usando archivos .bat (Más fácil)", level=2)
    add_formatted_paragraph(doc,
        "El método más sencillo para iniciar la aplicación consiste en hacer doble clic sobre el archivo ejecutar.bat "
        "(o EJECUTAR_APLICACION.bat para la versión premium) que se encuentra en la carpeta de la aplicación. Al "
        "ejecutar este archivo, se abrirá automáticamente una ventana de consola que mostrará información sobre el "
        "inicio del sistema, incluyendo las características de la versión premium como la paleta de colores moderna, "
        "gráficos optimizados, animaciones suaves, efectos glassmorphism y tipografías profesionales. Después de unos "
        "segundos, la aplicación se abrirá automáticamente en el navegador web predeterminado del equipo.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "4.2. Método 2: Usando comandos", level=2)
    add_formatted_paragraph(doc,
        "Los usuarios más avanzados pueden optar por iniciar la aplicación directamente desde la línea de comandos. "
        "Para ello, abra una ventana de PowerShell o CMD, navegue hasta la carpeta de la aplicación y ejecute el "
        "comando: \"streamlit run app.py\". Este método ofrece mayor control sobre el proceso de inicio y permite "
        "ver mensajes de depuración en tiempo real que pueden ser útiles para diagnosticar cualquier problema que "
        "pudiera presentarse durante la ejecución.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "4.3. Acceso desde el navegador", level=2)
    add_formatted_paragraph(doc,
        "Independientemente del método utilizado para iniciar la aplicación, una vez que el servidor Streamlit esté "
        "en funcionamiento, la aplicación será accesible a través de la dirección URL http://localhost:8501 en "
        "cualquier navegador web del equipo. Si la aplicación no se abre automáticamente, el usuario puede copiar "
        "esta dirección y pegarla en la barra de direcciones de su navegador preferido. Se recomienda utilizar "
        "Google Chrome o Microsoft Edge para obtener la mejor experiencia visual y de rendimiento. Para detener "
        "la aplicación en cualquier momento, basta con presionar la combinación de teclas Ctrl+C en la ventana "
        "de consola donde se está ejecutando el servidor.",
        font_size=11, first_line_indent=1.0)

    doc.add_page_break()

    # ================================================================
    # CAPÍTULO 5: ESTRUCTURA DE LA APLICACIÓN
    # ================================================================
    add_heading_styled(doc, "5. Estructura de la Aplicación", level=1)

    add_heading_styled(doc, "5.1. Arquitectura de archivos", level=2)
    add_formatted_paragraph(doc,
        "La aplicación GPP se compone de varios archivos organizados en una única carpeta denominada Aplicacion_GPP. "
        "Cada archivo cumple una función específica dentro del sistema. El archivo principal es app.py, que contiene "
        "todo el código fuente de la aplicación, incluyendo la lógica de negocio, la interfaz de usuario, los estilos "
        "CSS, las matrices de evaluación, los algoritmos de cálculo y las funciones de visualización. Los archivos "
        ".bat (ejecutar.bat, instalar.bat y EJECUTAR_APLICACION.bat) son scripts de automatización que facilitan "
        "la instalación y ejecución del sistema sin necesidad de conocimientos técnicos. Los archivos de texto "
        "(LEEME_PRIMERO.txt, INSTRUCCIONES_RAPIDAS.txt y COMO_ABRIR_HTML.txt) proporcionan documentación rápida "
        "y accesible directamente desde el explorador de archivos.",
        font_size=11, first_line_indent=1.0)

    add_image_with_caption(doc, os.path.join(IMG_DIR, "11_estructura_archivos.png"),
                           "Figura 2. Estructura de archivos de la aplicación GPP y su función.")

    add_heading_styled(doc, "5.2. Flujo de trabajo general", level=2)
    add_formatted_paragraph(doc,
        "El flujo de trabajo de la aplicación GPP sigue un proceso lógico y secuencial diseñado para guiar al usuario "
        "desde la instalación hasta la obtención de resultados. El proceso completo consta de cinco etapas principales: "
        "primero, la instalación de dependencias (solo necesaria una vez); segundo, la ejecución de la aplicación "
        "mediante los archivos .bat o por línea de comandos; tercero, la evaluación de los procesos administrativos y "
        "operativos mediante la calificación de cada variable en la escala de 1 a 5; cuarto, la consulta de los "
        "resultados generales y el análisis de prioridades a través de gráficos interactivos y tablas detalladas; y "
        "quinto, el guardado de la evaluación en formato JSON o la exportación de resultados a Excel para su uso en "
        "reportes formales. Las etapas tres a cinco pueden repetirse tantas veces como sea necesario, permitiendo "
        "al usuario realizar múltiples evaluaciones a lo largo del tiempo y comparar los avances obtenidos.",
        font_size=11, first_line_indent=1.0)

    add_image_with_caption(doc, os.path.join(IMG_DIR, "10_flujo_trabajo.png"),
                           "Figura 3. Diagrama del flujo de trabajo general de la aplicación GPP.")

    doc.add_page_break()

    # ================================================================
    # CAPÍTULO 6: PANTALLA DE INICIO
    # ================================================================
    add_heading_styled(doc, "6. Pantalla de Inicio", level=1)

    add_heading_styled(doc, "6.1. Bienvenida y guía de uso", level=2)
    add_formatted_paragraph(doc,
        "Al acceder a la aplicación por primera vez, el usuario es recibido por la pantalla de inicio, que presenta "
        "un diseño profesional y moderno con un banner de bienvenida en tonos corporativos (azul oscuro a verde teal). "
        "Esta pantalla constituye el punto de partida de la experiencia del usuario y proporciona toda la información "
        "necesaria para comprender el propósito de la herramienta y cómo utilizarla correctamente.",
        font_size=11, first_line_indent=1.0)

    add_formatted_paragraph(doc,
        "La sección de bienvenida incluye un instructivo completo que explica la naturaleza de la herramienta: es una "
        "aplicación diseñada para facilitar la autoevaluación de la gestión por procesos en establecimientos, "
        "identificando fortalezas y oportunidades de mejora. El instructivo detalla paso a paso cómo navegar por "
        "las diferentes secciones utilizando el menú lateral izquierdo, cómo realizar la evaluación de los procesos "
        "administrativos y operativos, y cómo interpretar los resultados obtenidos.",
        font_size=11, first_line_indent=1.0)

    add_image_with_caption(doc, os.path.join(IMG_DIR, "01_pantalla_inicio.png"),
                           "Figura 4. Pantalla de inicio de la aplicación GPP con el banner de bienvenida y la guía de uso.")

    add_formatted_paragraph(doc,
        "En el panel lateral izquierdo de la pantalla se encuentra el menú de navegación principal, que permite acceder "
        "a las seis secciones de la aplicación: Inicio, Proceso Administrativo, Proceso Operativo, Resultados Generales, "
        "Prioridades y Guardar/Cargar. Cada opción del menú está representada por un icono descriptivo y su nombre "
        "correspondiente, facilitando la identificación rápida de cada sección. La opción actualmente seleccionada se "
        "resalta con un color distintivo (verde teal) para indicar al usuario su ubicación dentro de la aplicación.",
        font_size=11, first_line_indent=1.0)

    add_image_with_caption(doc, os.path.join(IMG_DIR, "02_menu_navegacion.png"),
                           "Figura 5. Menú de navegación lateral con las seis secciones principales de la aplicación.",
                           width=Inches(2.5))

    add_heading_styled(doc, "6.2. Escala de calificación", level=2)
    add_formatted_paragraph(doc,
        "Uno de los elementos más importantes de la pantalla de inicio es la explicación de la escala de calificación "
        "utilizada en toda la evaluación. El sistema emplea una escala numérica de 1 a 5, donde cada valor tiene un "
        "significado específico y claramente definido. El valor 1 (No cumple) indica que el establecimiento no cumple "
        "en absoluto con el criterio evaluado; el valor 2 (Cumple parcialmente) señala un cumplimiento mínimo o "
        "incipiente del criterio; el valor 3 (Cumple de forma aceptable) refleja un nivel de cumplimiento moderado "
        "con oportunidades de mejora evidentes; el valor 4 (Cumple en gran medida) indica un nivel de cumplimiento "
        "alto con solo oportunidades menores de mejora; y el valor 5 (Cumple plenamente) significa que el criterio "
        "se cumple de forma completa y satisfactoria.",
        font_size=11, first_line_indent=1.0)

    add_image_with_caption(doc, os.path.join(IMG_DIR, "08_escala_calificacion.png"),
                           "Figura 6. Escala de calificación de 1 a 5 utilizada en la evaluación GPP.")

    add_heading_styled(doc, "6.3. Barra de progreso", level=2)
    add_formatted_paragraph(doc,
        "En la parte derecha de la pantalla de inicio se encuentra la tarjeta de progreso de la evaluación. Este "
        "indicador muestra en tiempo real el porcentaje de avance del usuario en la evaluación, calculado como la "
        "proporción de preguntas respondidas sobre el total de preguntas disponibles. La barra de progreso se actualiza "
        "automáticamente cada vez que el usuario califica una nueva variable, proporcionando una referencia visual "
        "inmediata sobre cuánto falta para completar la evaluación. Debajo del porcentaje, se muestra el conteo "
        "exacto de preguntas completadas versus el total, ofreciendo una referencia numérica precisa del avance.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "6.4. Niveles de prioridad", level=2)
    add_formatted_paragraph(doc,
        "También en la pantalla de inicio se presenta la tarjeta de niveles de prioridad, que explica al usuario "
        "cómo se clasifican los resultados de la evaluación. El sistema utiliza tres niveles de prioridad basados "
        "en el porcentaje de cumplimiento obtenido: Prioridad Baja (representada en color verde) para elementos con "
        "un cumplimiento igual o superior al 75%, lo que indica que estos aspectos se encuentran en buen estado y no "
        "requieren atención inmediata; Prioridad Media (representada en color amarillo) para elementos con un "
        "cumplimiento entre 60% y 74%, que deben ser atendidos en el mediano plazo; y Prioridad Alta (representada "
        "en color rojo) para elementos con un cumplimiento inferior al 60%, que requieren acción correctiva inmediata "
        "o a corto plazo.",
        font_size=11, first_line_indent=1.0)

    doc.add_page_break()

    # ================================================================
    # CAPÍTULO 7: PROCESO ADMINISTRATIVO
    # ================================================================
    add_heading_styled(doc, "7. Evaluación del Proceso Administrativo", level=1)

    add_formatted_paragraph(doc,
        "La sección de Evaluación del Proceso Administrativo constituye uno de los dos módulos principales de "
        "evaluación de la aplicación GPP. Para acceder a esta sección, el usuario debe seleccionar la opción "
        "\"Proceso Administrativo\" en el menú de navegación lateral. Esta sección evalúa cuatro aspectos "
        "fundamentales de la gestión administrativa del establecimiento: Planeación, Organización, Dirección y Control. "
        "Cada aspecto contiene múltiples elementos, y cada elemento presenta una o más variables específicas que deben "
        "ser calificadas individualmente.",
        font_size=11, first_line_indent=1.0)

    add_image_with_caption(doc, os.path.join(IMG_DIR, "03_proceso_administrativo.png"),
                           "Figura 7. Sección de evaluación del Proceso Administrativo con sus formularios de calificación.")

    add_heading_styled(doc, "7.1. Planeación", level=2)
    add_formatted_paragraph(doc,
        "El aspecto de Planeación evalúa la capacidad del establecimiento para anticipar y prepararse estratégicamente "
        "para el futuro. Se divide en dos elementos principales: el Análisis del contexto y la Existencia de un plan "
        "estratégico. El Análisis del contexto comprende tres variables que examinan si la empresa realiza análisis "
        "externo e interno como base para su planeación estratégica, y si aplica herramientas de análisis como la "
        "matriz DOFA, PESTEL, las 5 Fuerzas de Porter o Benchmarking. Por su parte, la Existencia de un plan "
        "estratégico evalúa cuatro variables relacionadas con la definición y comunicación de la misión, la visión, "
        "los objetivos estratégicos y la existencia de un documento formal que evidencie el plan estratégico a largo "
        "o mediano plazo. Este elemento tiene un peso relativo del 53.75% para el Análisis del contexto y 46.25% "
        "para el plan estratégico dentro del aspecto de Planeación.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "7.2. Organización", level=2)
    add_formatted_paragraph(doc,
        "El aspecto de Organización evalúa la estructura y los procesos internos del establecimiento. Se compone de "
        "cuatro elementos: la Existencia de una estructura organizativa (que evalúa si el establecimiento cuenta con "
        "un organigrama claramente definido, con un peso del 11.25%), la Departamentalización (que verifica si se "
        "tienen claramente identificados los procesos de la organización, con un peso del 35%), los Procesos "
        "documentados (que examina la existencia de perfiles de cargo, la documentación de procesos mediante mapas, "
        "flujogramas y procedimientos, y la caracterización de los procesos, con un peso del 33.75%), y el Ciclo PHVA "
        "(que evalúa la implementación del ciclo Planear-Hacer-Verificar-Actuar en el establecimiento, con un peso "
        "del 20%). Cada variable dentro de estos elementos debe ser calificada de 1 a 5 según el nivel de cumplimiento "
        "observado.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "7.3. Dirección", level=2)
    add_formatted_paragraph(doc,
        "El aspecto de Dirección se enfoca en evaluar las competencias de liderazgo, comunicación y cultura "
        "organizacional del establecimiento. Comprende tres elementos: Liderazgo organizacional (que evalúa la "
        "existencia de líderes proactivos y motivadores, con un peso del 35%), Canales de comunicación efectivos "
        "(que verifica la existencia de sistemas de comunicación formal tanto vertical como horizontal para el flujo "
        "de información, con un peso del 30%), y Cultura organizacional (que examina cinco variables relacionadas "
        "con el trabajo en equipo, la innovación, el reconocimiento de logros, el código de ética y el sentido de "
        "pertenencia del personal, con un peso del 35%). Este aspecto es fundamental para entender la dinámica humana "
        "y de liderazgo dentro de la organización.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "7.4. Control", level=2)
    add_formatted_paragraph(doc,
        "El aspecto de Control evalúa los mecanismos de seguimiento, medición y mejora implementados en el "
        "establecimiento. Se estructura en tres elementos: Sistemas de control (que evalúa el seguimiento periódico "
        "al cumplimiento de objetivos estratégicos y el uso de indicadores de gestión, con un peso del 45%), "
        "Auditorías internas (que verifica la realización periódica de auditorías para evaluar el cumplimiento de "
        "procesos, con un peso del 25%), y Acciones correctivas y preventivas (que examina la existencia de "
        "procedimientos para implementar acciones correctivas y preventivas, con un peso del 30%). Estos elementos "
        "son esenciales para garantizar que la organización aprende de sus errores y mejora continuamente.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "7.5. Cómo calificar las variables", level=2)
    add_formatted_paragraph(doc,
        "Para calificar cada variable, el usuario debe desplegar el elemento correspondiente haciendo clic en su "
        "título (indicado con el icono de chincheta). Dentro de cada elemento expandido, encontrará las preguntas "
        "o criterios de evaluación listados en la columna izquierda, y un menú desplegable (selectbox) en la columna "
        "derecha donde podrá seleccionar la calificación de 0 a 5. El valor 0 indica que la pregunta no ha sido "
        "respondida aún. Una vez calificadas todas las variables del Proceso Administrativo, el usuario debe hacer "
        "clic en el botón \"Guardar Evaluación del Proceso Administrativo\" ubicado en la parte inferior de la "
        "página. El sistema confirmará el guardado exitoso mediante un mensaje de éxito de color verde.",
        font_size=11, first_line_indent=1.0)

    doc.add_page_break()

    # ================================================================
    # CAPÍTULO 8: PROCESO OPERATIVO
    # ================================================================
    add_heading_styled(doc, "8. Evaluación del Proceso Operativo", level=1)

    add_formatted_paragraph(doc,
        "La sección de Evaluación del Proceso Operativo es el segundo módulo principal de evaluación de la "
        "aplicación GPP, al cual se accede seleccionando \"Proceso Operativo\" en el menú de navegación lateral. "
        "Esta sección evalúa tres aspectos fundamentales relacionados con las operaciones del establecimiento "
        "gastronómico: la Logística de Compras, la Gestión de Producción y la Logística Externa. La mecánica de "
        "evaluación es idéntica a la del Proceso Administrativo: el usuario califica cada variable en la escala "
        "de 1 a 5 y guarda los resultados al finalizar.",
        font_size=11, first_line_indent=1.0)

    add_image_with_caption(doc, os.path.join(IMG_DIR, "04_proceso_operativo.png"),
                           "Figura 8. Sección de evaluación del Proceso Operativo con formularios de calificación.")

    add_heading_styled(doc, "8.1. Logística de Compras", level=2)
    add_formatted_paragraph(doc,
        "El aspecto de Logística de Compras evalúa la eficiencia y organización del proceso de adquisición y "
        "almacenamiento de insumos. Se divide en dos elementos: la Logística de entrada (con un peso del 60%) que "
        "evalúa tres variables fundamentales (la organización en la planeación de compras considerando periodicidad, "
        "proveedores, cantidades e inventarios; el uso de herramientas de control como formatos, bases de datos o "
        "software; y las condiciones adecuadas de almacenamiento según el tipo de producto), y la Gestión de "
        "inventarios (con un peso del 40%) que evalúa dos variables (la implementación de un sistema de inventario "
        "como PEPS o ABC, y la automatización del inventario mediante software de gestión). Una logística de compras "
        "eficiente es fundamental para garantizar la disponibilidad y calidad de las materias primas.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "8.2. Gestión de Producción", level=2)
    add_formatted_paragraph(doc,
        "El aspecto de Gestión de Producción es el más extenso del Proceso Operativo y evalúa seis elementos "
        "fundamentales: Equipos (que examina la adecuación de los equipos al volumen de producción y las actividades "
        "de mantenimiento, con un peso del 15%); Infraestructura (que evalúa la ventilación, el diseño de cocina y "
        "la iluminación, con un peso del 15%); Distribución de los procesos o flujo (que verifica la identificación "
        "de áreas de trabajo y la eficiencia del flujo de trabajo, con un peso del 10%); Planeación de la producción "
        "(que evalúa cinco variables incluyendo niveles de producción, recetas estandarizadas, control de tiempos, "
        "optimización de recursos y manejo de desperdicios, con un peso del 25%); Materia prima (que verifica "
        "estándares de calidad y control de recepción, con un peso del 15%); y Control de la producción (que evalúa "
        "el control de calidad durante el proceso, el registro de datos y las acciones de mejora, con un peso del "
        "20%). En conjunto, estos elementos proporcionan una visión integral de la capacidad productiva del "
        "establecimiento.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "8.3. Logística Externa", level=2)
    add_formatted_paragraph(doc,
        "El aspecto de Logística Externa se centra en la interacción del establecimiento con sus clientes y evalúa "
        "un único elemento denominado Distribución, que tiene un peso del 100% dentro de este aspecto. Este elemento "
        "comprende cinco variables que examinan la agilidad y eficiencia del servicio al cliente, la capacitación del "
        "personal de atención, el seguimiento a la satisfacción del cliente, la gestión adecuada de quejas y reclamos, "
        "y la implementación de mejoras basadas en la retroalimentación de los clientes. La logística externa es "
        "fundamental para la fidelización de clientes y la reputación del establecimiento.",
        font_size=11, first_line_indent=1.0)

    doc.add_page_break()

    # ================================================================
    # CAPÍTULO 9: RESULTADOS GENERALES
    # ================================================================
    add_heading_styled(doc, "9. Resultados Generales", level=1)

    add_formatted_paragraph(doc,
        "La sección de Resultados Generales es donde la aplicación GPP transforma los datos de la evaluación en "
        "información visual significativa y procesable. Para acceder a esta sección, seleccione \"Resultados Generales\" "
        "en el menú de navegación lateral. Es importante haber completado al menos parcialmente las evaluaciones de "
        "los procesos administrativo y operativo antes de consultar los resultados; de lo contrario, el sistema "
        "mostrará un mensaje de advertencia indicando que no hay evaluaciones registradas.",
        font_size=11, first_line_indent=1.0)

    add_image_with_caption(doc, os.path.join(IMG_DIR, "05_resultados_generales.png"),
                           "Figura 9. Pantalla de Resultados Generales con indicador gauge, gráfico de barras y análisis radial.")

    add_heading_styled(doc, "9.1. Panel de métricas principales", level=2)
    add_formatted_paragraph(doc,
        "En la parte superior de la sección de resultados se presenta un panel con tres métricas clave dispuestas "
        "en columnas: el Cumplimiento General (expresado como porcentaje), la Calificación Promedio (expresada como "
        "valor sobre 5.0) y el Nivel de Prioridad (indicado con su icono de color y nombre correspondiente). Estas "
        "métricas proporcionan una visión instantánea del estado general del establecimiento, permitiendo al usuario "
        "comprender de un vistazo el resultado global de su evaluación. Las tarjetas de métricas incorporan efectos "
        "visuales interactivos como animaciones de hover y gradientes de color que mejoran la experiencia visual.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "9.2. Indicador de cumplimiento (Gauge)", level=2)
    add_formatted_paragraph(doc,
        "Debajo del panel de métricas se encuentra el indicador de tipo gauge o velocímetro, que representa "
        "gráficamente el nivel de cumplimiento general del establecimiento en una escala de 0% a 100%. El indicador "
        "utiliza tres zonas de color para facilitar la interpretación: la zona naranja (0-60%) indica un cumplimiento "
        "insuficiente que requiere acción inmediata, la zona amarilla (60-75%) señala un cumplimiento moderado que "
        "necesita atención, y la zona verde (75-100%) refleja un buen nivel de cumplimiento. Una línea de referencia "
        "amarilla marca la meta mínima del 75%, permitiendo al usuario visualizar rápidamente si su establecimiento "
        "alcanza o supera este umbral. El valor numérico del porcentaje se muestra prominentemente en el centro del "
        "indicador junto con la diferencia respecto a la meta del 75%.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "9.3. Gráfico de barras por área", level=2)
    add_formatted_paragraph(doc,
        "El gráfico de barras horizontales presenta el porcentaje de cumplimiento desglosado por cada aspecto "
        "evaluado, tanto del Proceso Administrativo (PA) como del Proceso Operativo (PO). Cada barra está coloreada "
        "según el nivel de prioridad correspondiente: verde para cumplimiento alto (75% o más), amarillo para "
        "cumplimiento medio (60-74%) y naranja para cumplimiento bajo (menos del 60%). El porcentaje exacto se "
        "muestra a la derecha de cada barra, facilitando la comparación precisa entre los diferentes aspectos. Este "
        "gráfico es especialmente útil para identificar rápidamente qué áreas del establecimiento presentan los "
        "mayores desafíos y cuáles son las más fortalecidas.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "9.4. Análisis radial multidimensional", level=2)
    add_formatted_paragraph(doc,
        "El gráfico radial (o gráfico de araña) ofrece una perspectiva multidimensional de los resultados, "
        "permitiendo visualizar simultáneamente el cumplimiento de todos los aspectos evaluados en un único diagrama. "
        "El gráfico presenta tres capas superpuestas: el área verde punteada que representa el nivel ideal del 100%, "
        "el área amarilla discontinua que marca la meta mínima del 75%, y el área azul oscuro sólido que muestra el "
        "cumplimiento actual del establecimiento. Los puntos donde el área de cumplimiento actual se acerca o supera "
        "la línea de meta mínima indican fortalezas, mientras que los puntos donde se aleja significativamente señalan "
        "áreas de oportunidad. Este tipo de visualización es particularmente valioso para presentaciones ejecutivas y "
        "comparaciones temporales.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "9.5. Tablas de resultados detallados", level=2)
    add_formatted_paragraph(doc,
        "En la parte inferior de la sección de resultados se encuentran las tablas de resultados detallados, "
        "organizadas en dos pestañas: una para el Proceso Administrativo y otra para el Proceso Operativo. Cada "
        "tabla muestra, para cada elemento evaluado dentro de cada aspecto, el promedio de calificación obtenido "
        "(en escala de 1 a 5), el porcentaje de cumplimiento correspondiente y la clasificación de prioridad "
        "(Alta, Media o Baja). Adicionalmente, al inicio de cada grupo de aspecto se muestra el porcentaje de "
        "cumplimiento consolidado del aspecto completo. Estas tablas complementan los gráficos proporcionando los "
        "datos numéricos exactos que sustentan las visualizaciones.",
        font_size=11, first_line_indent=1.0)

    doc.add_page_break()

    # ================================================================
    # CAPÍTULO 10: PRIORIDADES
    # ================================================================
    add_heading_styled(doc, "10. Análisis de Prioridades", level=1)

    add_formatted_paragraph(doc,
        "La sección de Análisis de Prioridades es una de las funcionalidades más valiosas de la aplicación GPP, "
        "ya que traduce los resultados numéricos de la evaluación en recomendaciones accionables para la mejora "
        "continua. Para acceder a esta sección, seleccione \"Prioridades\" en el menú de navegación lateral.",
        font_size=11, first_line_indent=1.0)

    add_image_with_caption(doc, os.path.join(IMG_DIR, "06_prioridades.png"),
                           "Figura 10. Sección de Análisis de Prioridades con gráfico de distribución y tarjetas de conteo.")

    add_heading_styled(doc, "10.1. Clasificación de prioridades", level=2)
    add_formatted_paragraph(doc,
        "Al ingresar a la sección de prioridades, lo primero que el usuario encontrará es un panel informativo "
        "detallado que explica los tres niveles de prioridad y su significado: la Prioridad Alta (indicada con el "
        "icono naranja) agrupa todos los elementos cuyo cumplimiento es inferior al 60% y que, por tanto, deben ser "
        "implementados o corregidos en el corto plazo con carácter de urgencia; la Prioridad Media (indicada con "
        "el icono amarillo) comprende los elementos con cumplimiento entre 60% y 74% que deben ser atendidos en "
        "el mediano plazo como parte de un plan de mejoramiento progresivo; y la Prioridad Baja (indicada con el "
        "icono verde) incluye los elementos con cumplimiento del 75% o superior que se encuentran en buen estado "
        "y no requieren atención inmediata, aunque siempre pueden beneficiarse de acciones de mantenimiento y "
        "mejora continua.",
        font_size=11, first_line_indent=1.0)

    add_formatted_paragraph(doc,
        "Debajo del panel informativo se presentan tres tarjetas que muestran el conteo de elementos clasificados "
        "en cada nivel de prioridad. Estas tarjetas permiten al usuario dimensionar rápidamente la proporción de "
        "elementos que requieren atención inmediata versus aquellos que se encuentran en condiciones satisfactorias. "
        "Un número elevado de elementos en prioridad alta indica que el establecimiento necesita un plan de mejora "
        "integral con acciones inmediatas, mientras que una predominancia de elementos en prioridad baja refleja "
        "una gestión por procesos sólida y bien implementada.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "10.2. Gráfico de distribución", level=2)
    add_formatted_paragraph(doc,
        "El gráfico de distribución de prioridades se presenta como un diagrama de dona (donut chart) que muestra "
        "visualmente la proporción de elementos en cada nivel de prioridad. En el centro del diagrama se indica el "
        "número total de elementos evaluados, mientras que cada segmento de la dona está coloreado según su nivel "
        "de prioridad (rojo para alta, amarillo para media y verde para baja) y muestra tanto la cantidad de "
        "elementos como el porcentaje que representan sobre el total. Este gráfico es especialmente útil para "
        "presentaciones ejecutivas ya que comunica de manera intuitiva y visualmente atractiva el estado general "
        "de la gestión por procesos del establecimiento.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "10.3. Tablas detalladas por nivel", level=2)
    add_formatted_paragraph(doc,
        "La sección de prioridades también incluye tres pestañas con tablas detalladas, una para cada nivel de "
        "prioridad. Cada tabla muestra información completa de los elementos clasificados en ese nivel, incluyendo "
        "el área a la que pertenecen (Proceso Administrativo o Proceso Operativo), el aspecto específico, el nombre "
        "del elemento, el porcentaje de cumplimiento obtenido y la clasificación de prioridad. Estas tablas permiten "
        "al usuario identificar exactamente cuáles son los elementos que requieren atención en cada nivel, "
        "facilitando la priorización de acciones de mejora y la elaboración de planes de trabajo específicos.",
        font_size=11, first_line_indent=1.0)

    doc.add_page_break()

    # ================================================================
    # CAPÍTULO 11: GUARDAR, CARGAR Y EXPORTAR
    # ================================================================
    add_heading_styled(doc, "11. Guardar, Cargar y Exportar", level=1)

    add_formatted_paragraph(doc,
        "La sección de Guardar/Cargar proporciona las funcionalidades necesarias para la persistencia y portabilidad "
        "de los datos de evaluación. Esta sección se divide en tres áreas funcionales: el guardado de evaluaciones, "
        "la carga de evaluaciones anteriores y la exportación de resultados a Excel.",
        font_size=11, first_line_indent=1.0)

    add_image_with_caption(doc, os.path.join(IMG_DIR, "07_guardar_cargar.png"),
                           "Figura 11. Sección de Guardar/Cargar con opciones de persistencia y exportación de datos.")

    add_heading_styled(doc, "11.1. Guardar evaluaciones (JSON)", level=2)
    add_formatted_paragraph(doc,
        "La función de guardado permite al usuario preservar la evaluación actual en un archivo de formato JSON "
        "(JavaScript Object Notation), que es un formato de intercambio de datos ligero, legible y ampliamente "
        "compatible. Al hacer clic en el botón \"Guardar Evaluación\", el sistema genera automáticamente un archivo "
        "con un nombre que incluye la fecha y hora exacta del guardado (por ejemplo, evaluacion_20251017_152909.json), "
        "lo que facilita la identificación y organización de múltiples evaluaciones realizadas en diferentes momentos. "
        "El archivo se guarda en la misma carpeta de la aplicación y contiene todas las calificaciones asignadas tanto "
        "para el Proceso Administrativo como para el Proceso Operativo, junto con la marca temporal del guardado. "
        "Se recomienda guardar la evaluación frecuentemente para evitar la pérdida de datos en caso de cierre "
        "inesperado de la aplicación.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "11.2. Cargar evaluaciones anteriores", level=2)
    add_formatted_paragraph(doc,
        "La función de carga permite al usuario recuperar evaluaciones previamente guardadas para consultarlas, "
        "continuarlas o compararlas con evaluaciones más recientes. El sistema muestra automáticamente una lista "
        "desplegable con todas las evaluaciones guardadas (archivos JSON) encontradas en la carpeta de la aplicación, "
        "identificadas por su fecha y hora de creación. Para cargar una evaluación, basta con seleccionarla en el "
        "menú desplegable y hacer clic en el botón \"Cargar Evaluación\". El sistema reemplazará las calificaciones "
        "actuales con las de la evaluación cargada y confirmará la operación con un mensaje de éxito. Esta "
        "funcionalidad es especialmente valiosa para realizar seguimientos periódicos de la gestión por procesos "
        "y evaluar la evolución del establecimiento a lo largo del tiempo.",
        font_size=11, first_line_indent=1.0)

    add_heading_styled(doc, "11.3. Exportar a Excel", level=2)
    add_formatted_paragraph(doc,
        "La función de exportación a Excel genera un archivo en formato XLSX con los resultados consolidados de la "
        "evaluación, que puede ser utilizado para la elaboración de reportes formales, presentaciones a directivos "
        "o archivamiento documental. El archivo exportado incluye una hoja de resumen con el área (Proceso "
        "Administrativo o Proceso Operativo), el aspecto evaluado y el porcentaje de cumplimiento correspondiente. "
        "Para exportar los resultados, el usuario debe hacer clic en el botón \"Exportar a Excel\" ubicado en la "
        "parte inferior de la sección. El archivo se genera automáticamente con un nombre que incluye la fecha y "
        "hora de la exportación (por ejemplo, resultados_gpp_20251017_153045.xlsx) y se guarda en la carpeta de "
        "la aplicación. Este archivo puede abrirse con Microsoft Excel, LibreOffice Calc o cualquier aplicación "
        "compatible con el formato XLSX.",
        font_size=11, first_line_indent=1.0)

    doc.add_page_break()

    # ================================================================
    # CAPÍTULO 12: VERSIÓN HTML
    # ================================================================
    add_heading_styled(doc, "12. Versión HTML Alternativa", level=1)

    add_formatted_paragraph(doc,
        "Además de la versión principal basada en Streamlit, la aplicación GPP incluye una versión alternativa en "
        "formato HTML (archivo app_gpp.html) que ofrece una experiencia simplificada pero completamente funcional "
        "que no requiere la instalación de Python ni de ninguna dependencia adicional. Esta versión es ideal para "
        "presentaciones, demostraciones, compartir con usuarios sin conocimientos técnicos o para situaciones en "
        "las que no sea posible instalar el software necesario para la versión completa.",
        font_size=11, first_line_indent=1.0)

    add_formatted_paragraph(doc,
        "Para utilizar la versión HTML, existen tres métodos de acceso igualmente válidos. El primer método y más "
        "sencillo consiste simplemente en hacer doble clic sobre el archivo app_gpp.html, lo que abrirá la "
        "aplicación automáticamente en el navegador predeterminado del equipo. El segundo método consiste en abrir "
        "manualmente el navegador web, presionar Ctrl+O para abrir el diálogo de apertura de archivos, y navegar "
        "hasta la ubicación del archivo HTML. El tercer método consiste en arrastrar el archivo app_gpp.html "
        "directamente a la ventana del navegador web abierto. La versión HTML incluye navegación por pestañas, "
        "formularios de evaluación, secciones de resultados con gráficos interactivos (generados con Chart.js) y "
        "un diseño responsive que se adapta a diferentes tamaños de pantalla.",
        font_size=11, first_line_indent=1.0)

    add_formatted_paragraph(doc,
        "Es importante considerar que la versión HTML tiene algunas limitaciones en comparación con la versión "
        "Streamlit: la persistencia de datos es limitada (los datos no se guardan automáticamente al cerrar el "
        "navegador), las funciones de exportación a Excel no están disponibles, y los cálculos estadísticos son "
        "más básicos. Sin embargo, para una visualización rápida del diseño y concepto de la aplicación, o para "
        "compartir la herramienta con personas que no puedan instalar el entorno Python, la versión HTML es una "
        "alternativa excelente. La versión HTML requiere conexión a Internet para cargar las fuentes de Google "
        "y la biblioteca Chart.js desde sus respectivos CDN.",
        font_size=11, first_line_indent=1.0)

    doc.add_page_break()

    # ================================================================
    # CAPÍTULO 13: SOLUCIÓN DE PROBLEMAS
    # ================================================================
    add_heading_styled(doc, "13. Solución de Problemas", level=1)

    add_formatted_paragraph(doc,
        "A continuación se presentan las situaciones problemáticas más comunes que pueden presentarse durante la "
        "instalación o uso de la aplicación GPP, junto con sus respectivas soluciones:",
        font_size=11, first_line_indent=1.0)

    # Tabla de problemas
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    for i, h in enumerate(["Problema", "Solución"]):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "172139")

    problems = [
        ("Python no está instalado o no se reconoce el comando",
         "Descargue e instale Python desde https://www.python.org/. Durante la instalación, asegúrese de marcar la opción \"Add Python to PATH\"."),
        ("Error al instalar dependencias",
         "Ejecute manualmente: py -m pip install -r requirements.txt --upgrade. Verifique su conexión a Internet."),
        ("La aplicación no se abre en el navegador",
         "Copie la URL http://localhost:8501 y péguela manualmente en la barra de direcciones de su navegador."),
        ("Los gráficos no se muestran correctamente",
         "Actualice su navegador a la última versión. Se recomienda Google Chrome o Microsoft Edge."),
        ("Error \"ModuleNotFoundError\"",
         "Ejecute: py verificar_instalacion.py para identificar los módulos faltantes y luego reinstale con: py -m pip install -r requirements.txt"),
        ("La evaluación no se guarda",
         "Verifique que tiene permisos de escritura en la carpeta de la aplicación. Intente ejecutar como administrador."),
        ("La aplicación se cierra inesperadamente",
         "No cierre la ventana de consola (CMD) mientras usa la aplicación. Presione Ctrl+C solo cuando desee detenerla."),
    ]

    for row_idx, (problem, solution) in enumerate(problems):
        table.rows[row_idx + 1].cells[0].text = problem
        table.rows[row_idx + 1].cells[1].text = solution
        for cell in table.rows[row_idx + 1].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()

    doc.add_page_break()

    # ================================================================
    # CAPÍTULO 14: RECOMENDACIONES DE USO
    # ================================================================
    add_heading_styled(doc, "14. Recomendaciones de Uso", level=1)

    add_formatted_paragraph(doc,
        "Para obtener el máximo beneficio de la Herramienta de Evaluación GPP, se recomienda seguir las siguientes "
        "buenas prácticas durante su utilización:",
        font_size=11, first_line_indent=1.0)

    recommendations = [
        ("Sea honesto y objetivo en sus calificaciones",
         "La utilidad de la herramienta depende directamente de la veracidad de las respuestas. Califique cada variable "
         "de manera objetiva, basándose en evidencia real y no en percepciones subjetivas o aspiraciones. Una evaluación "
         "sincera, aunque muestre resultados desfavorables, es infinitamente más valiosa que una evaluación inflada "
         "que no refleje la realidad del establecimiento."),
        ("Guarde su evaluación frecuentemente",
         "Durante el proceso de evaluación, es recomendable guardar el progreso periódicamente utilizando la función "
         "de guardado. Esto previene la pérdida de datos en caso de cierre inesperado del navegador o la aplicación."),
        ("Priorice las acciones de mejora en elementos de prioridad alta",
         "Los elementos clasificados como Prioridad Alta (cumplimiento inferior al 60%) deben ser atendidos de manera "
         "prioritaria. Desarrolle planes de acción específicos y medibles para cada uno de estos elementos, con "
         "responsables, plazos y recursos asignados."),
        ("Realice evaluaciones periódicas",
         "Se recomienda realizar la evaluación al menos cada trimestre o semestre para monitorear la evolución de la "
         "gestión por procesos. Compare los resultados de evaluaciones sucesivas para identificar tendencias de mejora "
         "o deterioro en cada área."),
        ("Exporte a Excel para reportes formales",
         "Utilice la función de exportación a Excel cuando necesite incluir los resultados en informes de gestión, "
         "presentaciones a directivos o documentación de auditorías. El formato Excel facilita la integración de los "
         "datos con otros sistemas de información."),
        ("Involucre al equipo de trabajo",
         "Aunque la evaluación puede ser realizada por una sola persona, se recomienda involucrar a diferentes miembros "
         "del equipo en el proceso para obtener perspectivas diversas y una evaluación más completa y equilibrada."),
    ]

    for title, desc in recommendations:
        p_title = doc.add_paragraph()
        run = p_title.add_run(f"  {title}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(25, 88, 85)
        add_formatted_paragraph(doc, desc, font_size=11, first_line_indent=1.0)

    doc.add_page_break()

    # ================================================================
    # CAPÍTULO 15: GLOSARIO
    # ================================================================
    add_heading_styled(doc, "15. Glosario de Términos", level=1)

    glossary = [
        ("Ciclo PHVA", "Metodología de gestión compuesta por cuatro etapas: Planear, Hacer, Verificar y Actuar. Se utiliza para la mejora continua de procesos."),
        ("Dashboard", "Panel de control visual que presenta indicadores clave de rendimiento y métricas relevantes de manera consolidada."),
        ("DOFA", "Herramienta de análisis estratégico que identifica Debilidades, Oportunidades, Fortalezas y Amenazas de una organización."),
        ("Glassmorphism", "Tendencia de diseño visual que utiliza transparencias, difuminados y bordes sutiles para crear efectos de vidrio esmerilado."),
        ("GPP", "Gestión Por Procesos. Enfoque administrativo que organiza las actividades de una empresa en procesos interrelacionados."),
        ("JSON", "JavaScript Object Notation. Formato de intercambio de datos ligero y legible utilizado para almacenar las evaluaciones."),
        ("PESTEL", "Herramienta de análisis del entorno que examina factores Políticos, Económicos, Sociales, Tecnológicos, Ecológicos y Legales."),
        ("Plotly", "Biblioteca de Python para la creación de gráficos interactivos de alta calidad."),
        ("PEPS", "Primeras Entradas, Primeras Salidas. Método de valuación de inventarios que establece que los primeros productos en ingresar son los primeros en ser utilizados."),
        ("Selectbox", "Elemento de interfaz de usuario que permite seleccionar un valor de una lista desplegable de opciones."),
        ("Streamlit", "Framework de Python para la creación de aplicaciones web interactivas orientadas a datos."),
        ("XLSX", "Formato de archivo de Microsoft Excel utilizado para hojas de cálculo. Es el formato de exportación de resultados de la aplicación."),
    ]

    table = doc.add_table(rows=len(glossary) + 1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(["Término", "Definición"]):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "172139")

    for idx, (term, definition) in enumerate(glossary):
        table.rows[idx + 1].cells[0].text = term
        table.rows[idx + 1].cells[1].text = definition
        for cell in table.rows[idx + 1].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
        # Bold term
        for p in table.rows[idx + 1].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # ================================================================
    # PIE DE PÁGINA FINAL
    # ================================================================
    p_final = doc.add_paragraph()
    p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_final.add_run("━" * 50)
    run.font.color.rgb = RGBColor(0, 174, 172)
    run.font.size = Pt(10)

    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_footer.add_run("Herramienta de Evaluación GPP - Versión Premium 2025\n"
                            "Sistema Profesional de Gestión Por Procesos\n"
                            "CEITTO - Colmayor | Proyecto Gastrocolma\n"
                            f"Manual generado el {datetime.now().strftime('%d de %B de %Y')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 116, 139)
    run.italic = True

    # Guardar documento
    doc_path = os.path.join(OUTPUT_DIR, "Manual_de_Usuario_GPP_Gastrocolma.docx")
    doc.save(doc_path)
    print(f"  ✅ Documento Word guardado: {doc_path}")
    return doc_path


# ============================================================
# PARTE 3: CONVERSIÓN A PDF
# ============================================================

def convertir_a_pdf(doc_path):
    """Convierte el documento Word a PDF"""
    print("\n📄 Convirtiendo a PDF...")
    try:
        from docx2pdf import convert
        pdf_path = doc_path.replace('.docx', '.pdf')
        convert(doc_path, pdf_path)
        print(f"  ✅ PDF generado: {pdf_path}")
        return pdf_path
    except Exception as e:
        print(f"  ⚠️ No se pudo convertir automáticamente a PDF: {e}")
        print("  Intentando método alternativo...")
        try:
            import subprocess
            # Intentar con LibreOffice
            result = subprocess.run([
                'soffice', '--headless', '--convert-to', 'pdf',
                '--outdir', OUTPUT_DIR, doc_path
            ], capture_output=True, text=True, timeout=60)
            if result.returncode == 0:
                pdf_path = doc_path.replace('.docx', '.pdf')
                print(f"  ✅ PDF generado con LibreOffice: {pdf_path}")
                return pdf_path
        except:
            pass

        # Intentar con COM de Word
        try:
            import comtypes.client
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            doc = word.Documents.Open(doc_path)
            pdf_path = doc_path.replace('.docx', '.pdf')
            doc.SaveAs(pdf_path, FileFormat=17)
            doc.Close()
            word.Quit()
            print(f"  ✅ PDF generado con Microsoft Word: {pdf_path}")
            return pdf_path
        except:
            pass

        print("  ❌ No se pudo generar el PDF automáticamente.")
        print("  Para generar el PDF manualmente:")
        print("     1. Abra el archivo .docx con Microsoft Word")
        print("     2. Seleccione Archivo > Guardar como > PDF")
        return None


# ============================================================
# EJECUCIÓN PRINCIPAL
# ============================================================

if __name__ == "__main__":
    print("=" * 60)
    print("  GENERADOR DE MANUAL DE USUARIO - APLICACIÓN GPP")
    print("  Proyecto Gastrocolma - CEITTO Colmayor")
    print("=" * 60)
    print(f"\n📂 Directorio de salida: {OUTPUT_DIR}")

    # Paso 1: Generar imágenes
    generar_todas_las_imagenes()

    # Paso 2: Generar documento Word
    doc_path = crear_documento_word()

    # Paso 3: Convertir a PDF
    pdf_path = convertir_a_pdf(doc_path)

    print("\n" + "=" * 60)
    print("  ✅ PROCESO COMPLETADO")
    print("=" * 60)
    print(f"\n📄 Word: {doc_path}")
    if pdf_path:
        print(f"📄 PDF:  {pdf_path}")
    print(f"📸 Imágenes: {IMG_DIR}")
    print(f"\n📂 Todo guardado en: {OUTPUT_DIR}")
    print("=" * 60)
