# -*- coding: utf-8 -*-
"""
Generador del Manual Técnico de la Aplicación GPP
Salida: Word y PDF en la misma carpeta
"""

import os
import sys
import io
from datetime import datetime

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

OUTPUT_DIR = r"C:\Users\ASUS\OneDrive\Desktop\Manual de uso Gastrocolma"
DOC_PATH = os.path.join(OUTPUT_DIR, "Manual_Tecnico_GPP_Gastrocolma.docx")


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Consolas' if level == 3 and any(c in text for c in ['()', 'def ', '.py']) else 'Calibri'
        if level == 1:
            run.font.color.rgb = RGBColor(23, 33, 57)
        elif level == 2:
            run.font.color.rgb = RGBColor(25, 88, 85)
        else:
            run.font.color.rgb = RGBColor(16, 129, 129)
    return h


def add_para(doc, text, indent=1.0, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(51, 65, 85)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(indent)
    return p


def add_code_block(doc, code, caption=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 41, 59)
    # Fondo gris claro
    pf = p.paragraph_format
    pf.space_after = Pt(8)
    if caption:
        cap = doc.add_paragraph()
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(8)
        cap_run.font.color.rgb = RGBColor(100, 116, 139)
        cap.paragraph_format.space_after = Pt(12)


def crear_manual_tecnico():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ----- PORTADA -----
    for _ in range(3):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("MANUAL TECNICO")
    r.bold = True
    r.font.size = Pt(32)
    r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(23, 33, 57)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Aplicacion de Evaluacion de Gestion Por Procesos (GPP)")
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(25, 88, 85)

    doc.add_paragraph()
    ln = doc.add_paragraph()
    ln.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ln.add_run("_" * 50).font.color.rgb = RGBColor(0, 174, 172)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info.add_run("Arquitectura, modelos de datos, algoritmos y guia de desarrollo\nProyecto Gastrocolma - CEITTO Colmayor")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Version: Premium 2025  |  Fecha: " + datetime.now().strftime("%Y-%m-%d")).font.size = Pt(10)

    doc.add_page_break()

    # ----- 1. RESUMEN EJECUTIVO -----
    add_heading_styled(doc, "1. Resumen ejecutivo", level=1)
    add_para(doc,
        "La aplicacion GPP es una aplicacion web monolítica desarrollada en Python que expone una interfaz de usuario "
        "mediante el framework Streamlit. Replica la logica de negocio y los criterios de evaluacion del archivo Excel "
        "\"Herramienta de evaluacion GPP_0 (7).xlsm\" en un entorno web con graficos interactivos (Plotly), persistencia "
        "en JSON y exportacion a Excel. Este manual describe la arquitectura tecnica, los modelos de datos, los "
        "algoritmos de calculo y las pautas para mantenimiento y extension.")

    doc.add_page_break()

    # ----- 2. ARQUITECTURA Y STACK TECNOLÓGICO -----
    add_heading_styled(doc, "2. Arquitectura y stack tecnologico", level=1)

    add_heading_styled(doc, "2.1. Diagrama de capas", level=2)
    add_para(doc,
        "La aplicacion sigue un modelo de tres capas logicas dentro de un unico proceso: capa de presentacion (Streamlit "
        "+ HTML/CSS inyectado), capa de logica de negocio (funciones Python en app.py que calculan prioridades y "
        "porcentajes) y capa de datos (session state en memoria + archivos JSON y XLSX en disco). No existe base de "
        "datos relacional; la persistencia se realiza mediante archivos en el sistema de archivos del servidor.")

    add_heading_styled(doc, "2.2. Tecnologias utilizadas", level=2)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'
    headers = ["Componente", "Version / Uso", "Rol en la aplicacion"]
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(c, "172139")

    rows = [
        ("Python", "3.8+ (recomendado 3.10+)", "Lenguaje base y entorno de ejecucion"),
        ("Streamlit", "1.x", "Framework web: UI, widgets, routing por sidebar"),
        ("Pandas", "2.x", "DataFrames para tablas y exportacion Excel"),
        ("Plotly", "5.x", "Graficos interactivos (barras, gauge, radial, pie)"),
        ("NumPy", "1.x", "Calculos numericos (promedios, conversiones)"),
        ("OpenPyXL", "3.x", "Motor de escritura de archivos .xlsx"),
    ]
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            table.rows[i + 1].cells[j].text = cell_text
            for p in table.rows[i + 1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    add_para(doc,
        "El archivo requirements.txt (cuando existe) debe listar estas dependencias con versiones compatibles. "
        "El script verificar_instalacion.py comprueba la presencia de streamlit, pandas, openpyxl, plotly y numpy.")

    add_heading_styled(doc, "2.3. Flujo de ejecucion", level=2)
    add_para(doc,
        "Al ejecutar \"streamlit run app.py\", Streamlit inicia un servidor HTTP (por defecto en el puerto 8501), "
        "carga app.py y ejecuta el script de arriba a abajo. En cada interaccion del usuario (cambio de pagina, "
        "seleccion en un widget, clic en boton), Streamlit re-ejecuta el script completo. El estado se mantiene "
        "gracias a st.session_state, que persiste entre reruns dentro de la misma sesion del navegador. Al cerrar "
        "el navegador o el servidor, el estado en memoria se pierde; solo persisten los archivos JSON y XLSX "
        "generados por el usuario.")

    doc.add_page_break()

    # ----- 3. ESTRUCTURA DEL PROYECTO -----
    add_heading_styled(doc, "3. Estructura del proyecto", level=1)

    add_heading_styled(doc, "3.1. Archivos principales", level=2)
    table = doc.add_table(rows=10, cols=2)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = "Archivo"
    table.rows[0].cells[1].text = "Descripcion tecnica"
    for c in table.rows[0].cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(c, "172139")

    files_desc = [
        ("app.py", "Punto de entrada. Contiene configuracion de pagina, CSS, matrices, pesos, funciones de calculo y todas las pantallas (main, mostrar_inicio, mostrar_evaluacion_pa/po, calcular_resultados, mostrar_resultados_generales, mostrar_prioridades, mostrar_guardar_cargar)."),
        ("verificar_instalacion.py", "Script independiente que importa streamlit, pandas, openpyxl, plotly y numpy e imprime OK/ERROR por modulo."),
        ("ejecutar.bat", "Invoca streamlit run app.py en la carpeta actual."),
        ("EJECUTAR_APLICACION.bat", "Variante que activa venv si existe y luego streamlit run app.py."),
        ("instalar.bat", "Ejecuta py -m pip install -r requirements.txt (y upgrade pip)."),
        ("requirements.txt", "Lista de paquetes pip (streamlit, pandas, plotly, numpy, openpyxl) con versiones."),
        ("evaluacion_*.json", "Generados por guardar_evaluacion(): contienen fecha, evaluaciones_pa y evaluaciones_po."),
        ("resultados_gpp_*.xlsx", "Generados al exportar a Excel: hoja 'Resumen' con Area, Aspecto, Cumplimiento (%)."),
        ("Herramienta de evaluacion GPP_0 (7).xlsm", "Excel de referencia; la logica de evaluacion de app.py replica este modelo."),
    ]
    for i, (f, d) in enumerate(files_desc):
        table.rows[i + 1].cells[0].text = f
        table.rows[i + 1].cells[1].text = d
        for cell in table.rows[i + 1].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    add_heading_styled(doc, "3.2. Dependencias entre modulos", level=2)
    add_para(doc,
        "app.py no importa modulos propios del proyecto; todo el codigo esta en un unico archivo. Los imports son: "
        "streamlit (st), pandas (pd), plotly.express (px), plotly.graph_objects (go), plotly.subplots (make_subplots), "
        "numpy (np), json, os y datetime. verificar_instalacion.py no depende de app.py; solo hace __import__ de los "
        "cinco paquetes listados.")

    doc.add_page_break()

    # ----- 4. MODELOS DE DATOS -----
    add_heading_styled(doc, "4. Modelos de datos", level=1)

    add_heading_styled(doc, "4.1. Matrices de evaluacion (MATRIZ_PA, MATRIZ_PO)", level=2)
    add_para(doc,
        "Las matrices definen la estructura de la evaluacion: aspectos, elementos y variables (preguntas). Son diccionarios "
        "anidados de tres niveles: aspecto -> elemento -> lista de cadenas (preguntas). MATRIZ_PA cubre el Proceso "
        "Administrativo: PLANEACION (2 elementos), ORGANIZACION (4 elementos), DIRECCION (3 elementos), CONTROL "
        "(3 elementos). MATRIZ_PO cubre el Proceso Operativo: LOGISTICA DE COMPRAS (2 elementos), GESTION DE "
        "PRODUCCION (6 elementos), LOGISTICA EXTERNA (1 elemento). Cada variable se califica con un valor entero "
        "de 0 a 5 (0 = sin responder).")

    add_heading_styled(doc, "4.2. Pesos relativos (PESOS_PA, PESOS_PO)", level=2)
    add_para(doc,
        "PESOS_PA y PESOS_PO son diccionarios aspecto -> elemento -> numero (porcentaje). Los pesos suman 100% dentro "
        "de cada aspecto. Se usan para documentar la importancia relativa de cada elemento; en la implementacion actual "
        "los calculos de cumplimiento por aspecto se realizan como promedios simples de los promedios de los elementos "
        "(no se aplica ponderacion explícita por peso en el codigo de calcular_resultados). Los pesos estan alineados "
        "con el diseno del Excel de referencia.")

    add_heading_styled(doc, "4.3. Session state (st.session_state)", level=2)
    add_para(doc,
        "evaluaciones_pa y evaluaciones_po son diccionarios clave-valor. Clave: cadena con formato "
        "\"PA_<aspecto>_<elemento>_<indice>\" o \"PO_<aspecto>_<elemento>_<indice>\". Valor: entero 0-5. "
        "resultados_calculados es un booleano que se pone en True al guardar una evaluacion PA o PO. La inicializacion "
        "se hace en inicializar_session_state() al inicio de main().")

    add_heading_styled(doc, "4.4. Formato JSON de persistencia", level=2)
    add_code_block(doc,
        '{\n  "fecha": "YYYY-MM-DD HH:MM:SS",\n  "evaluaciones_pa": { "PA_...": 0-5, ... },\n  "evaluaciones_po": { "PO_...": 0-5, ... }\n}',
        "Estructura del archivo generado por guardar_evaluacion().")

    add_heading_styled(doc, "4.5. Estructura de resultados (calcular_resultados)", level=2)
    add_para(doc,
        "La funcion calcular_resultados() devuelve un diccionario con claves 'pa', 'po' y 'general'. 'pa' y 'po' son "
        "diccionarios aspecto -> (elemento -> dict con promedio, porcentaje, prioridad, tipo_prioridad, peso, "
        "calificaciones) mas claves promedio_aspecto y porcentaje_aspecto por aspecto. 'general' contiene 'promedio' "
        "y 'porcentaje' del establecimiento (promedio de todos los promedios de aspecto).")

    doc.add_page_break()

    # ----- 5. ALGORITMOS Y LÓGICA DE NEGOCIO -----
    add_heading_styled(doc, "5. Algoritmos y logica de negocio", level=1)

    add_heading_styled(doc, "5.1. Calculo de prioridad", level=2)
    add_code_block(doc,
        "def calcular_prioridad(porcentaje):\n    if porcentaje >= 75: return \"BAJA\", \"success\"\n    elif porcentaje >= 60: return \"MEDIA\", \"warning\"\n    else: return \"ALTA\", \"error\"",
        "Umbrales: BAJA >= 75%, MEDIA 60-74%, ALTA < 60%.")

    add_heading_styled(doc, "5.2. Cumplimiento por elemento y porcentaje", level=2)
    add_code_block(doc,
        "def calcular_cumplimiento_elemento(calificaciones):\n    return np.mean(calificaciones) if calificaciones else 0\n\ndef calcular_porcentaje_cumplimiento(promedio):\n    return (promedio / 5) * 100",
        "Las calificaciones 0 se excluyen en el bucle que construye la lista de calificaciones por elemento (solo se anaden valores > 0).")

    add_heading_styled(doc, "5.3. Flujo de calcular_resultados()", level=2)
    add_para(doc,
        "Para cada aspecto en MATRIZ_PA/MATRIZ_PO se recorren sus elementos y variables. Por cada variable se obtiene "
        "la clave en session state (PA_/PO_ + aspecto + elemento + indice) y, si la calificacion es > 0, se anade a "
        "una lista. Con esa lista se calcula el promedio del elemento, luego el porcentaje (promedio/5*100), la "
        "prioridad (calcular_prioridad) y se guarda el peso. El promedio del aspecto es la media de los promedios de "
        "sus elementos. El resultado general es la media de todos los promedios de aspecto (PA y PO).")

    doc.add_page_break()

    # ----- 6. INTERFAZ DE USUARIO (STREAMLIT) -----
    add_heading_styled(doc, "6. Interfaz de usuario (Streamlit)", level=1)

    add_heading_styled(doc, "6.1. Configuracion de pagina", level=2)
    add_code_block(doc,
        "st.set_page_config(page_title=\"Evaluacion GPP\", page_icon=\"...\", layout=\"wide\", initial_sidebar_state=\"expanded\")",
        "Pagina en modo wide; sidebar visible al cargar.")

    add_heading_styled(doc, "6.2. Estilos CSS", level=2)
    add_para(doc,
        "Se inyecta un bloque st.markdown(..., unsafe_allow_html=True) con CSS que define variables (--primary, "
        "--secondary, --accent, colores de prioridad, gradientes, sombras), estilos para .main-header, .sub-header, "
        ".metric-card, .stButton, sidebar, expanders, progress bar, dataframes, tabs y footer. Se importan fuentes "
        "Google (Montserrat, Roboto, Outfit). El diseno es corporativo con paleta azul/verde teal.")

    add_heading_styled(doc, "6.3. Navegacion", level=2)
    add_para(doc,
        "st.sidebar.radio con seis opciones determina la pagina actual. main() despacha a una de: mostrar_inicio(), "
        "mostrar_evaluacion_pa(), mostrar_evaluacion_po(), mostrar_resultados_generales(), mostrar_prioridades(), "
        "mostrar_guardar_cargar(). No hay URLs ni rutas; todo es condicional en el rerun del script.")

    add_heading_styled(doc, "6.4. Widgets por pantalla", level=2)
    add_para(doc,
        "Inicio: markdown HTML, st.progress(progreso). PA/PO: st.expander por elemento, st.selectbox por variable "
        "(options=[0,1,2,3,4,5], key unico por variable), st.button para guardar. Resultados: st.metric (3 columnas), "
        "st.plotly_chart (gauge, barras, radial), st.tabs + st.dataframe. Prioridades: metricas HTML, pie chart "
        "Plotly, st.tabs con dataframes. Guardar/Cargar: st.button Guardar, st.selectbox de archivos JSON, "
        "st.button Cargar, st.button Exportar a Excel.")

    add_heading_styled(doc, "6.5. Graficos (Plotly)", level=2)
    add_para(doc,
        "Se utilizan go.Figure y go.Indicator (gauge), go.Bar (barras horizontales), go.Scatterpolar (radial), "
        "go.Pie (dona). Colores y titulos se alinean con la paleta CSS (prioridad alta/media/baja). Los graficos "
        "son interactivos (hover, zoom) y se insertan con st.plotly_chart(..., use_container_width=True).")

    doc.add_page_break()

    # ----- 7. PERSISTENCIA Y EXPORTACIÓN -----
    add_heading_styled(doc, "7. Persistencia y exportacion", level=1)

    add_heading_styled(doc, "7.1. Guardar evaluacion (JSON)", level=2)
    add_para(doc,
        "guardar_evaluacion() construye un dict con fecha (datetime.now()), evaluaciones_pa y evaluaciones_po, "
        "y lo escribe con json.dump en un archivo evaluacion_YYYYMMDD_HHMMSS.json en el directorio de app.py "
        "(os.path.dirname(__file__)). Codificacion UTF-8, ensure_ascii=False, indent=2.")

    add_heading_styled(doc, "7.2. Cargar evaluacion", level=2)
    add_para(doc,
        "cargar_evaluacion(archivo) lee el JSON, asigna datos['evaluaciones_pa'] y datos['evaluaciones_po'] a "
        "st.session_state y pone resultados_calculados en True. Luego se llama st.rerun() para refrescar la UI "
        "con los datos cargados.")

    add_heading_styled(doc, "7.3. Exportar a Excel", level=2)
    add_para(doc,
        "Se usa pd.ExcelWriter con engine openpyxl. Se crea una hoja 'Resumen' con columnas Area, Aspecto, "
        "Cumplimiento (%). Se rellena con los porcentajes por aspecto de resultados['pa'] y resultados['po']. "
        "El archivo se guarda como resultados_gpp_YYYYMMDD_HHMMSS.xlsx en el directorio de app.py.")

    doc.add_page_break()

    # ----- 8. DESPLIEGUE Y EJECUCIÓN -----
    add_heading_styled(doc, "8. Despliegue y ejecucion", level=1)

    add_heading_styled(doc, "8.1. Entorno de ejecucion", level=2)
    add_para(doc,
        "Requisito: Python 3.8 o superior instalado y accesible desde la linea de comandos (py o python). Las "
        "dependencias se instalan con pip (pip install -r requirements.txt o instalar.bat). No se requiere servidor "
        "web externo ni base de datos; Streamlit actua como servidor HTTP y sirve la aplicacion en http://localhost:8501.")

    add_heading_styled(doc, "8.2. Comandos de arranque", level=2)
    add_code_block(doc,
        "streamlit run app.py\n# Opcional: activar venv antes (Windows)\n# call venv\\Scripts\\activate.bat",
        "El directorio de trabajo debe ser la carpeta que contiene app.py.")

    add_heading_styled(doc, "8.3. Puertos y red", level=2)
    add_para(doc,
        "Por defecto Streamlit usa el puerto 8501. Si el puerto esta ocupado, se puede usar --server.port: "
        "streamlit run app.py --server.port 8502. La aplicacion es local; para acceso desde otras maquinas se "
        "requeriria exponer el puerto o usar un tunel/reverse proxy.")

    doc.add_page_break()

    # ----- 9. MANTENIMIENTO Y EXTENSIÓN -----
    add_heading_styled(doc, "9. Mantenimiento y extension", level=1)

    add_heading_styled(doc, "9.1. Modificar preguntas o pesos", level=2)
    add_para(doc,
        "Editar las constantes MATRIZ_PA, MATRIZ_PO, PESOS_PA y PESOS_PO en app.py. Aumentar o reducir el numero "
        "de variables en un elemento implica que las claves de session state (indice i en el bucle) sigan siendo "
        "consistentes; si se eliminan variables, las claves antiguas pueden quedar huérfanas en evaluaciones "
        "guardadas, por lo que conviene documentar cambios de version en el formato JSON.")

    add_heading_styled(doc, "9.2. Anadir una nueva pagina", level=2)
    add_para(doc,
        "Anadir una opcion al radio del sidebar, una rama en el if/elif de main() que llame a una nueva funcion "
        "(por ejemplo mostrar_nueva_seccion()), e implementar esa funcion con los widgets necesarios. Si la nueva "
        "seccion usa resultados calculados, llamar a calcular_resultados() dentro de la funcion.")

    add_heading_styled(doc, "9.3. Cambiar umbrales de prioridad", level=2)
    add_para(doc,
        "Modificar la funcion calcular_prioridad(porcentaje). Los umbrales actuales son 75 y 60; cualquier cambio "
        "afecta a resultados, graficos y tablas de prioridades de forma automatica.")

    add_heading_styled(doc, "9.4. Pruebas recomendadas", level=2)
    add_para(doc,
        "Tras cambios en matrices o algoritmos: (1) Ejecutar verificar_instalacion.py. (2) Abrir la app, completar "
        "al menos un aspecto PA y uno PO, revisar Resultados y Prioridades. (3) Guardar evaluacion, cerrar y volver "
        "a abrir, cargar el JSON y comprobar que los valores y graficos coinciden. (4) Exportar a Excel y verificar "
        "la hoja Resumen.")

    doc.add_page_break()

    # ----- 10. REFERENCIAS Y GLOSARIO TÉCNICO -----
    add_heading_styled(doc, "10. Referencias y glosario tecnico", level=1)

    refs = [
        ("Streamlit", "https://docs.streamlit.io/ - Documentacion oficial del framework."),
        ("Plotly Python", "https://plotly.com/python/ - Graficos interactivos."),
        ("Pandas", "https://pandas.pydata.org/docs/ - DataFrames y Excel."),
        ("Session state", "st.session_state en Streamlit persiste datos entre reruns en la misma sesion del navegador."),
        ("Rerun", "Cada interaccion del usuario provoca la re-ejecucion completa del script app.py."),
    ]
    table = doc.add_table(rows=len(refs) + 1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = "Termino / Recurso"
    table.rows[0].cells[1].text = "Descripcion"
    for c in table.rows[0].cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(c, "172139")
    for i, (term, desc) in enumerate(refs):
        table.rows[i + 1].cells[0].text = term
        table.rows[i + 1].cells[1].text = desc
        for cell in table.rows[i + 1].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = foot.add_run("Manual Tecnico GPP - Proyecto Gastrocolma | CEITTO Colmayor | " +
                     datetime.now().strftime("%Y-%m-%d"))
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(100, 116, 139)
    r.italic = True

    doc.save(DOC_PATH)
    print("Documento Word guardado: " + DOC_PATH)
    return DOC_PATH


def convertir_a_pdf(doc_path):
    try:
        from docx2pdf import convert
        pdf_path = doc_path.replace('.docx', '.pdf')
        convert(doc_path, pdf_path)
        print("PDF generado: " + pdf_path)
        return pdf_path
    except Exception as e:
        print("No se pudo generar PDF automaticamente: " + str(e))
        return None


if __name__ == "__main__":
    print("=" * 60)
    print("  MANUAL TECNICO - APLICACION GPP (Gastrocolma)")
    print("=" * 60)
    doc_path = crear_manual_tecnico()
    pdf_path = convertir_a_pdf(doc_path)
    print("=" * 60)
    print("  Completado. Word: " + doc_path)
    if pdf_path:
        print("  PDF: " + pdf_path)
    print("=" * 60)
